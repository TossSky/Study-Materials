# -*- coding: utf-8 -*-
"""Генерация отчёта по курсовой работе ТВиМС (вариант 9, Log-Laplace).
Содержит пункты 1–5 курсовой:
  1) исследование распределения;
  2) знакомство с Python и статистики выборки;
  3) эмпирическая функция распределения;
  4) гистограмма;
  5) описание параметров теоретической функции распределения.

Запуск:  python generate_report.py
Выход:   Курсовая_ТВиМС_Вариант9.docx
"""

import json
import math
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "figures"

with open(ROOT / "numbers.json", "r", encoding="utf-8") as f:
    R = json.load(f)


def ru(x, n=4):
    """Форматирование числа с запятой в качестве десятичного разделителя."""
    return f"{x:.{n}f}".replace(".", ",")


# ---------- Глобальные стили (как в отчётах АСВТ) ----------
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def add_centered(text, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.bold = bold
    r.font.size = Pt(size); r.font.name = "Times New Roman"
    return p


def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(14); r.font.name = "Times New Roman"
    return p


def add_subheading(text):
    add_blank()
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.name = "Times New Roman"; r.font.size = Pt(14)
    return p


def add_body(text):
    return doc.add_paragraph(text)


def add_listing(code, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(caption); r.font.size = Pt(14); r.font.name = "Times New Roman"
    table = doc.add_table(rows=1, cols=1); table.style = "Table Grid"
    cell = table.cell(0, 0); cell.text = ""
    for i, line in enumerate(code.split("\n")):
        pp = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        pp.paragraph_format.first_line_indent = Cm(0)
        pp.paragraph_format.line_spacing = 1.0
        pp.paragraph_format.space_after = Pt(0)
        pp.paragraph_format.space_before = Pt(0)
        pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = pp.add_run(line); rr.font.name = "Courier New"; rr.font.size = Pt(10)


def add_table_label(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.font.size = Pt(14); r.font.name = "Times New Roman"


def add_fig_label(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.font.size = Pt(14); r.font.name = "Times New Roman"
    r.italic = True


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(h); r.font.name = "Times New Roman"
        r.font.size = Pt(12); r.bold = True
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]; c.text = ""
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(v)); r.font.name = "Times New Roman"; r.font.size = Pt(12)
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = w


def add_image(path, width_inch=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width_inch))


# ========== ТИТУЛЬНАЯ СТРАНИЦА ==========
add_centered("Министерство науки и высшего образования Российской Федерации")
add_centered("Санкт-Петербургский политехнический университет Петра Великого")
add_blank()
add_centered("Институт компьютерных наук и кибербезопасности")
add_centered("Высшая школа кибербезопасности")
for _ in range(4):
    add_blank()
add_centered("КУРСОВАЯ РАБОТА", bold=True, size=16)
add_centered("«Исследование выборки случайной величины и проверка")
add_centered("статистических гипотез»")
add_blank()
add_centered("по дисциплине")
add_centered("«Теория вероятностей и математическая статистика»")
add_blank()
add_centered("Вариант 9 (распределение Log-Laplace)")
for _ in range(4):
    add_blank()

for label in [
    "Выполнил",
    "студент гр. 5151003/40001\t\t\t\t\t\t\tТоцкий В.",
    "",
    "Преподаватель\t\t\t\t\t\t\t\t\tПахомова О. А.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(label); r.font.size = Pt(14); r.font.name = "Times New Roman"

for _ in range(3):
    add_blank()
add_centered("Санкт-Петербург – 2026")
doc.add_page_break()


# ========== 1. ЦЕЛЬ РАБОТЫ ==========
add_heading("1. Цель работы")
add_blank()
add_body(
    "Цель работы — провести статистическое исследование выборки X, подчиняющейся "
    "распределению Log-Laplace (вариант 9). В отчёте выполнены первые пять "
    "пунктов задания: обзор распределения и его прикладного назначения; "
    "реализация на языке Python функций вычисления выборочных статистик и "
    "эмпирических характеристик распределения без обращения к готовым "
    "реализациям; построение эмпирической функции распределения и гистограммы "
    "на подвыборках разных размеров; запись аналитического вида теоретической "
    "функции распределения Log-Laplace и анализ влияния её параметров."
)
add_body(
    "Исходные данные — файл var_9_loglaplace.csv с n = {n} вещественными "
    "наблюдениями. Расчёты выполнены в Jupyter Notebook на Python 3. NumPy и "
    "Matplotlib использованы только для визуализации; статистические "
    "процедуры (среднее, дисперсия, моменты, ЭФР, гистограмма, CDF), которые "
    "по заданию требуется реализовать самостоятельно, написаны вручную."
    .format(n=R["n"])
)
doc.add_page_break()


# ========== 2. ИССЛЕДОВАНИЕ РАСПРЕДЕЛЕНИЯ ==========
add_heading("2. Исследование распределения Log-Laplace")
add_blank()

add_subheading("2.1. История появления")
add_blank()
add_body(
    "Распределение Лапласа (двустороннее экспоненциальное) введено П.-С. Лапласом "
    "в трактате 1774 года «Mémoire sur la probabilité des causes par les "
    "événements» [1] как закон распределения ошибок наблюдений с плотностью, "
    "пропорциональной exp(−|x|/b). По отношению к нормальному закону Лапласово "
    "распределение характеризуется более тяжёлыми хвостами и экспоненциальным, "
    "а не квадратичным убыванием плотности."
)
add_body(
    "Распределение Log-Laplace получается экспоненциальным преобразованием: "
    "если ln X имеет распределение Лапласа с параметрами μ и b, то X называют "
    "логарифмически-Лапласовским. Систематическое изложение свойств Log-Laplace "
    "и связанных асимметричных обобщений приведено в монографии С. Котца, "
    "Т. Козубовского и К. Подгурского «The Laplace Distribution and "
    "Generalizations» [2]. У. Рид в работе [3] показал, что Log-Laplace "
    "(в его обозначениях — двойное распределение Парето) возникает как "
    "стационарный закон геометрического броуновского движения, остановленного "
    "в случайный экспоненциально распределённый момент; это объясняет, почему "
    "Log-Laplace эмпирически хорошо описывает данные, формирующиеся как "
    "результат мультипликативных случайных воздействий — размеры файлов, "
    "городов, доходов и пр."
)

add_subheading("2.2. Назначение и применение в научных исследованиях")
add_blank()
add_body(
    "Log-Laplace применяется для моделирования положительных величин, у которых "
    "логарифм подчиняется симметричному распределению с тяжёлыми хвостами. "
    "Типовые области применения:"
)
add_body(
    "— финансовая математика: лог-доходности активов на коротких временных "
    "горизонтах. Эмпирическая плотность лог-доходностей, как правило, имеет "
    "более острую вершину и более тяжёлые хвосты, чем нормальное "
    "приближение, — Лапласовские модели согласуются с этими наблюдениями "
    "лучше гауссовских [2];"
)
add_body(
    "— масштабные явления с мультипликативной природой: распределения "
    "размеров файлов, городов, доходов и других величин, формирующихся как "
    "результат мультипликативных случайных воздействий [3]."
)

add_subheading("2.3. Применение в задачах информационной безопасности")
add_blank()
add_body(
    "Сетевой трафик в реальных системах не описывается гауссовским законом: "
    "для длительностей соединений, размеров запросов и интервалов между "
    "событиями характерны тяжёлые хвосты. В работе М. Кровеллы и А. Беставроса "
    "[4] на наборе HTTP-трасс показана статистическая самоподобность и "
    "тяжелохвостовая природа объёмов передаваемых файлов и длительностей "
    "сессий. Это базовый эмпирический факт, на который опираются современные "
    "системы обнаружения сетевых аномалий: при гауссовском допущении частота "
    "ложных срабатываний оказывается недопустимо высокой."
)
add_body(
    "Log-Laplace в этом контексте занимает место параметрической модели для "
    "величин, чей логарифм симметрично распределён вокруг медианы. К таким "
    "задачам относятся: оценка распределения времён выполнения "
    "криптографических операций (анализ временных каналов утечки, "
    "timing side-channels), где центральная масса задержек узкая, а "
    "редкие выбросы — двусторонние; и моделирование размеров пакетов и "
    "объёмов лог-событий при настройке порогов обнаружения."
)

doc.add_page_break()


# ========== 3. ЗНАКОМСТВО С PYTHON. СТАТИСТИКИ ВЫБОРКИ ==========
add_heading("3. Знакомство с Python и статистики выборки")
add_blank()

add_subheading("3.1. Среда выполнения")
add_blank()
add_body(
    "Расчёты выполнены на Python 3.14 в среде Jupyter Notebook. Используются "
    "стандартная библиотека (модули csv, math, random) и пакеты numpy и "
    "matplotlib — последние только для отрисовки графиков. Считывание выборки "
    "из файла var_9_loglaplace.csv выполнено модулем csv:"
)
add_listing(
    "import csv\n"
    "sample = []\n"
    "with open('var_9_loglaplace.csv', 'r', encoding='utf-8') as f:\n"
    "    for row in csv.reader(f):\n"
    "        if row:\n"
    "            sample.append(float(row[0]))",
    "Листинг 1 — Считывание выборки",
)

add_subheading("3.2. Выборочные статистики")
add_blank()
add_body(
    "Все статистики вычислены вручную, без обращения к готовым реализациям "
    "(numpy.mean, statistics.median и т. п.). Соответствующие функции собраны "
    "в модуле compute.py и приведены в Приложении А. Используются следующие "
    "определения:"
)
add_body(
    "1. Сумма элементов: S = Σ xᵢ. 2. Выборочное среднее: x̄ = S/n. "
    "3. Медиана: серединный элемент вариационного ряда (для чётного n — "
    "среднее арифметическое двух центральных). 4. Мода: значение с "
    "максимальной частотой; для непрерывной выборки определена после "
    "округления (см. примечание). 5. Размах: R = xₘₐₓ − xₘᵢₙ. "
    "6. Смещённая (выборочная) дисперсия: D = (1/n) Σ (xᵢ − x̄)². "
    "7. Несмещённая дисперсия: s² = (1/(n−1)) Σ (xᵢ − x̄)². "
    "8. Выборочный начальный момент порядка k: νₖ = (1/n) Σ xᵢᵏ. "
    "9. Выборочный центральный момент порядка k: μₖ = (1/n) Σ (xᵢ − x̄)ᵏ."
)
add_blank()
add_table_label("Таблица 1 — Числовые значения выборочных статистик")
add_table(
    ["Статистика", "Обозначение", "Значение"],
    [
        ["Объём выборки", "n", f"{R['n']}"],
        ["Сумма элементов", "S", ru(R['sum'], 6)],
        ["Выборочное среднее", "x̄", ru(R['mean'], 6)],
        ["Медиана", "Me", ru(R['median'], 6)],
        ["Мода (округление до 0,01)", "Mo",
         f"{ru(R['mode_values'][0], 2)} (частота {R['mode_count']})"],
        ["Минимум", "x_min", ru(R['min'], 6)],
        ["Максимум", "x_max", ru(R['max'], 6)],
        ["Размах", "R", ru(R['range'], 6)],
        ["Смещённая дисперсия", "D", ru(R['var_biased'], 6)],
        ["Несмещённая дисперсия", "s²", ru(R['var_unbiased'], 6)],
        ["Выборочный начальный момент 3-го порядка", "ν₃",
         ru(R['moment_initial_3'], 6)],
        ["Выборочный начальный момент 4-го порядка", "ν₄",
         ru(R['moment_initial_4'], 6)],
        ["Выборочный центральный момент 3-го порядка", "μ₃",
         ru(R['moment_central_3'], 6)],
        ["Выборочный центральный момент 4-го порядка", "μ₄",
         ru(R['moment_central_4'], 6)],
    ],
)
add_blank()
add_body(
    "Примечание о моде. Все 300 наблюдений в выборке попарно различны "
    "(непрерывная случайная величина), поэтому в исходном виде моды у выборки "
    "нет — каждое значение встречается ровно один раз. Для получения "
    "практически интерпретируемой моды значения предварительно округлены до "
    "сотых; в этом представлении наибольшую частоту имеет точка "
    f"{ru(R['mode_values'][0], 2)}, что согласуется с медианой "
    f"({ru(R['median'], 4)}) и подтверждает близость к симметричному "
    "одномодальному распределению с центром около 7,17."
)
add_body(
    "Анализ значений: x̄ ≈ Me, что характерно для распределений, симметричных "
    "относительно центра в логарифмической шкале. Малая величина центрального "
    "момента третьего порядка (μ₃ ≈ 6,4·10⁻³ при разбросе значений около 1,5) "
    "подтверждает практическую симметричность данных. Отношение μ₄/D² ≈ "
    f"{ru(R['moment_central_4'] / R['var_biased']**2, 2)} существенно превышает "
    "значение 3, характерное для нормального закона, что свидетельствует о "
    "наличии тяжёлых хвостов и согласуется с гипотезой Log-Laplace."
)
doc.add_page_break()


# ========== 4. ЭМПИРИЧЕСКАЯ ФУНКЦИЯ РАСПРЕДЕЛЕНИЯ ==========
add_heading("4. Эмпирическая функция распределения")
add_blank()

add_subheading("4.1. Определение и алгоритм построения")
add_blank()
add_body(
    "Эмпирической функцией распределения выборки X = (x₁, …, xₙ) называется "
    "функция F_n(x), равная доле элементов выборки, не превосходящих x:"
)
add_body(
    "    F_n(x) = (1/n) · |{ i : xᵢ ≤ x }|.    (1)"
)
add_body(
    "Свойства F_n: 1) F_n — неубывающая ступенчатая функция; "
    "2) F_n(x) → 0 при x → −∞ и F_n(x) → 1 при x → +∞; "
    "3) скачки происходят в точках вариационного ряда x_(1) ≤ … ≤ x_(n), "
    "величина скачка в точке x_(k) равна (числу совпадающих с ней значений)/n; "
    "для непрерывной выборки все скачки равны 1/n. По теореме Гливенко — "
    "Кантелли F_n(x) сходится равномерно к теоретической функции "
    "распределения F(x) при n → ∞."
)
add_body(
    "Алгоритм построения. (1) Отсортировать выборку по возрастанию. "
    "(2) Сформировать значения F_n в точках x_(k): F_n(x_(k)) = k / n. "
    "(3) На интервале [x_(k), x_(k+1)) функция постоянна и равна k/n; "
    "при x < x_(1) — равна 0, при x ≥ x_(n) — равна 1."
)
add_body(
    "Реализация на Python без использования готовых процедур приведена в "
    "Приложении А (функция ecdf). Графики построены matplotlib.pyplot.step "
    "с параметром where='post'."
)

add_subheading("4.2. Графики ЭФР для подвыборок")
add_blank()
add_body(
    "Из исходной выборки сделаны случайные подвыборки объёмов 10, 100 и 200 "
    "(без возвращения, генератор numpy.random.default_rng с фиксированным "
    "зерном для воспроизводимости). Графики ЭФР приведены ниже."
)
for k in (10, 100, 200):
    add_image(FIG / f"ecdf_{k}.png", width_inch=5.6)
    add_fig_label(f"Рис. {1 if k==10 else 2 if k==100 else 3}. ЭФР подвыборки n = {k}")
add_image(FIG / "ecdf_full.png", width_inch=5.6)
add_fig_label("Рис. 4. ЭФР полной выборки n = 300")

add_subheading("4.3. Выводы по ЭФР")
add_blank()
add_body(
    "1. Форма ЭФР. Уже при n = 100 ЭФР приобретает гладкую S-образную форму с "
    "выраженным «изгибом» вблизи x ≈ 7,17. Эта точка близка к выборочной "
    f"медиане ({ru(R['median'], 4)}) и к значению exp(μ̂) ≈ "
    f"{ru(math.exp(R['mu_hat_mean_logs']), 4)}, где μ̂ = "
    f"{ru(R['mu_hat_mean_logs'], 4)} — среднее логарифмов выборки."
)
add_body(
    "2. Соответствие теоретическому виду. Полученный график визуально "
    "согласуется с CDF Log-Laplace (см. раздел 6): резкий рост в окрестности "
    "x = exp(μ) и медленный «выход на единицу» справа за счёт тяжёлого хвоста."
)
add_body(
    "3. Предположения о параметрах. По выборке получены грубые оценки: "
    f"μ̂ ≈ {ru(R['mu_hat_mean_logs'], 4)} (среднее логарифмов), "
    f"b̂ ≈ {ru(R['b_hat_from_var'], 4)} (из соотношения Var(ln X) = 2b² по "
    "несмещённой дисперсии логарифмов). Полученные значения согласуются с "
    "положением и крутизной перегиба ЭФР."
)
doc.add_page_break()


# ========== 5. ГИСТОГРАММА ==========
add_heading("5. Гистограмма")
add_blank()

add_subheading("5.1. Определение и алгоритм построения")
add_blank()
add_body(
    "Гистограммой выборки X называется кусочно-постоянная функция, "
    "построенная по разбиению области значений выборки на k непересекающихся "
    "интервалов (карманов) [c₀, c₁), [c₁, c₂), …, [c_{k−1}, c_k]. Для "
    "i-го интервала ширины Δᵢ = cᵢ − cᵢ₋₁, в который попало nᵢ наблюдений, "
    "значение гистограммы плотности равно:"
)
add_body(
    "    h_i = nᵢ / (n · Δᵢ).    (2)"
)
add_body(
    "При равной ширине интервалов Δᵢ ≡ Δ = (xₘₐₓ − xₘᵢₙ)/k. В пределе при "
    "n → ∞ и Δ → 0 (с условием nΔ → ∞) гистограмма сходится к плотности "
    "теоретического распределения f(x); таким образом, гистограмма — "
    "состоятельная оценка плотности."
)
add_body(
    "Алгоритм построения. (1) Зафиксировать xₘᵢₙ и xₘₐₓ — границы выборки. "
    "(2) Выбрать число интервалов k. В работе использована формула Стёрджеса "
    "k = round(1 + log₂ n), дающая для n = 300 значение k = 9; для подвыборок "
    "k = 4 (n = 10), k = 8 (n = 100), k = 9 (n = 200). "
    "(3) Задать границы cᵢ = xₘᵢₙ + i·Δ. "
    "(4) Подсчитать nᵢ — число элементов выборки в каждом интервале (для "
    "последнего интервала включается верхняя граница). "
    "(5) Построить столбцы высоты h_i = nᵢ/(n·Δ)."
)
add_body(
    "Реализация — функция histogram() в Приложении А; matplotlib.pyplot.bar "
    "использован только для отрисовки готовых столбцов."
)

add_subheading("5.2. Графики гистограмм для подвыборок")
add_blank()
for k in (10, 100, 200):
    add_image(FIG / f"hist_{k}.png", width_inch=5.6)
    add_fig_label(f"Рис. {5 if k==10 else 6 if k==100 else 7}. Гистограмма подвыборки n = {k}")
add_image(FIG / "hist_full.png", width_inch=5.6)
add_fig_label("Рис. 8. Гистограмма полной выборки n = 300")

add_subheading("5.3. Выводы по гистограмме")
add_blank()
add_body(
    "1. Форма гистограммы. На полной выборке гистограмма имеет одиночный "
    "максимум в кармане [7,02; 7,19) — 38 % наблюдений — и спад в обе "
    "стороны. Спад асимметричен в исходной шкале X: слева гистограмма "
    "доходит до xₘᵢₙ = 6,68 (расстояние от моды ≈ 0,49), справа — до "
    "xₘₐₓ = 8,20 (≈ 1,03), причём в правом хвосте при x > 7,7 находится "
    "5 наблюдений из 300 против 8 в левом хвосте при x < 6,85. Эта "
    "асимметрия согласуется с плотностью Log-Laplace: она симметрична "
    "относительно ln x = μ, но множитель 1/x в (3) растягивает её вправо "
    "в исходной шкале X."
)
add_body(
    "2. Зависимость от размера подвыборки. При n = 10 гистограмма крайне "
    "неинформативна: число пустых интервалов сопоставимо с числом "
    "заполненных, оценка плотности зашумлена. При n = 100 уже видна "
    "одномодальная структура. При n = 200 и тем более на полной выборке "
    "форма устойчива. Это иллюстрирует свойство гистограммы как "
    "состоятельной оценки плотности — улучшение с ростом n."
)
add_body(
    "3. Предположения о параметрах. Положение максимума гистограммы "
    "соответствует exp(μ̂); ширина «основания» столбцов с заметной "
    "плотностью (примерно от 6,9 до 7,5, т. е. размах около 0,6) "
    "согласуется с малым значением b̂ ≈ 0,018: при таких параметрах квантили "
    "уровней 0,05 и 0,95 для Log-Laplace равны соответственно "
    "exp(μ̂ + b̂·ln(0,1)) ≈ 6,89 и exp(μ̂ − b̂·ln(0,1)) ≈ 7,48, что хорошо "
    "соотносится с наблюдаемым диапазоном на гистограмме."
)
doc.add_page_break()


# ========== 6. ОПИСАНИЕ ПАРАМЕТРОВ РАСПРЕДЕЛЕНИЯ ==========
add_heading("6. Описание параметров распределения Log-Laplace")
add_blank()

add_subheading("6.1. Аналитический вид распределения")
add_blank()
add_body(
    "Случайная величина X имеет распределение Log-Laplace с параметром "
    "положения μ ∈ ℝ и параметром масштаба b > 0, если её логарифм "
    "Y = ln X распределён по закону Лапласа Laplace(μ, b). Плотность "
    "распределения:"
)
add_body(
    "    f(x; μ, b) = (1/(2 b x)) · exp(−|ln x − μ|/b),  x > 0,    (3)"
)
add_body(
    "где параметр μ — медиана распределения Y = ln X (так что exp(μ) — "
    "медиана X), а b — масштабный параметр Лапласовского закона. Функция "
    "распределения получается интегрированием (3) и записывается кусочно-"
    "аналитически:"
)
add_body(
    "    F(x; μ, b) = ½ · exp((ln x − μ)/b),                  0 < x ≤ exp(μ);"
)
add_body(
    "    F(x; μ, b) = 1 − ½ · exp(−(ln x − μ)/b),             x > exp(μ).    (4)"
)
add_body(
    "Эквивалентное компактное представление: F(x; μ, b) = ½ + ½ · sgn(ln x − μ) · "
    "[1 − exp(−|ln x − μ|/b)]."
)

add_subheading("6.2. Описание параметров")
add_blank()
add_table_label("Таблица 2 — Параметры распределения Log-Laplace")
add_table(
    ["Параметр", "Тип", "Область", "Содержательная интерпретация"],
    [
        ["μ", "положение", "ℝ",
         "математическое ожидание и медиана ln X; exp(μ) — медиана X"],
        ["b", "масштаб", "(0, +∞)",
         "параметр масштаба для ln X; Var(ln X) = 2b²"],
    ],
    col_widths=[Cm(2.0), Cm(2.6), Cm(2.5), Cm(8.5)],
)
add_blank()
add_body(
    "Существование моментов. Момент порядка k ∈ ℝ: E[X^k] = exp(kμ) / "
    "(1 − k²b²) при |k|·b < 1; иначе момент не существует. В частности, "
    "среднее E[X] = exp(μ) / (1 − b²) при b < 1, дисперсия конечна при "
    "b < 1/2. Тяжесть хвоста управляется параметром b: при малых b "
    "распределение сосредоточено в окрестности exp(μ), при b → 1/4 "
    "обращается в бесконечность четвёртый момент, при b → 1/2 — дисперсия."
)

add_subheading("6.3. Реализация теоретической функции распределения")
add_blank()
add_listing(
    "import math\n"
    "def loglaplace_cdf(x, mu, b):\n"
    "    if x <= 0:\n"
    "        return 0.0\n"
    "    z = math.log(x) - mu\n"
    "    if z <= 0:\n"
    "        return 0.5 * math.exp(z / b)\n"
    "    return 1.0 - 0.5 * math.exp(-z / b)",
    "Листинг 2 — Функция вычисления F(x; μ, b) для Log-Laplace",
)

add_subheading("6.4. Графики при различных значениях параметров")
add_blank()
add_body(
    "На рис. 9 приведены графики F(x; μ, b) для нескольких сочетаний "
    "параметров. Видно, что параметр μ определяет горизонтальное «положение» "
    "графика, а параметр b — крутизну центральной ступени и скорость выхода "
    "хвостов на асимптоты 0 и 1."
)
add_image(FIG / "loglaplace_examples.png", width_inch=5.8)
add_fig_label("Рис. 9. Теоретическая функция распределения Log-Laplace при "
              "различных сочетаниях параметров")
add_blank()
add_body(
    "Для количественной оценки роли каждого из параметров проведён парный "
    "анализ: при фиксации одного параметра второй принимает два сильно "
    "различающихся значения (рис. 10)."
)
add_image(FIG / "loglaplace_param_influence.png", width_inch=6.4)
add_fig_label("Рис. 10. Влияние параметров μ (слева) и b (справа) на "
              "теоретическую функцию распределения")

add_subheading("6.5. Выводы о влиянии параметров")
add_blank()
add_body(
    "1. Параметр μ — параметр сдвига в логарифмической шкале. При его "
    "увеличении на Δμ график F(x; μ, b) смещается вправо в исходных "
    "координатах в exp(Δμ) раз (мультипликативный сдвиг по оси x), "
    "сохраняя форму. Так, увеличение μ с 1,50 до 2,20 (Δμ = 0,70) сдвигает "
    "медиану exp(μ) с 4,48 до 9,03 — в exp(0,70) ≈ 2,01 раза."
)
add_body(
    "2. Параметр b — параметр масштаба в логарифмической шкале. При "
    "увеличении b центральный участок графика становится более пологим, "
    "хвосты — более медленно сходящимися к 0 и 1. Положение медианы exp(μ) "
    "не меняется (все кривые при одинаковом μ пересекаются в точке "
    "(exp(μ), 0,5)). На правой панели рис. 10 кривая b = 0,01 близка к "
    "ступенчатой функции единичного скачка в точке exp(1,97) ≈ 7,17, тогда "
    "как кривая b = 0,10 имеет заметный разброс в диапазоне x ∈ [5; 12]."
)
add_body(
    "3. Связь с выборкой. Грубая оценка b̂ ≈ 0,018 объясняет наблюдаемую "
    "узость гистограммы (раздел 5.3) и крутизну центральной части ЭФР "
    "(раздел 4.3): малое b означает узкое распределение ln X вокруг μ и, "
    "как следствие, концентрацию X вокруг exp(μ)."
)
add_blank()
add_body(
    "На рис. 11 совмещены ЭФР полной выборки и теоретическая функция "
    "распределения с подставленными грубыми оценками параметров "
    f"(μ̂ = {ru(R['mu_hat_mean_logs'], 4)}, b̂ = {ru(R['b_hat_from_var'], 4)}). "
    "Кривые визуально близки на всём диапазоне, расхождение наибольшее на "
    "правом хвосте при x > 7,8."
)
add_image(FIG / "ecdf_vs_theory.png", width_inch=5.8)
add_fig_label("Рис. 11. Совмещение ЭФР и теоретической F(x; μ̂, b̂)")
doc.add_page_break()


# ========== СПИСОК ЛИТЕРАТУРЫ ==========
add_heading("Список использованных источников")
add_blank()
refs = [
    "Laplace P.-S. Mémoire sur la probabilité des causes par les événements // "
    "Mémoires de l'Académie royale des sciences présentés par divers savants. — "
    "1774. — Vol. 6. — P. 621–656.",
    "Kotz S., Kozubowski T. J., Podgórski K. The Laplace Distribution and "
    "Generalizations: A Revisit with Applications to Communications, Economics, "
    "Engineering and Finance. — Boston: Birkhäuser, 2001. — 349 p. — "
    "ISBN 978-0-8176-4166-5.",
    "Reed W. J. The Pareto, Zipf and other power laws // Economics Letters. — "
    "2001. — Vol. 74, № 1. — P. 15–19. — DOI: 10.1016/S0165-1765(01)00524-9.",
    "Crovella M. E., Bestavros A. Self-similarity in World Wide Web traffic: "
    "evidence and possible causes // IEEE/ACM Transactions on Networking. — "
    "1997. — Vol. 5, № 6. — P. 835–846. — DOI: 10.1109/90.650143.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f"{i}. {ref}")
    p.paragraph_format.first_line_indent = Cm(0)
doc.add_page_break()


# ========== ПРИЛОЖЕНИЕ А ==========
add_heading("Приложение А. Листинги собственных реализаций")
add_blank()
add_listing(
    "def s_sum(x):\n"
    "    s = 0.0\n"
    "    for v in x:\n"
    "        s += v\n"
    "    return s\n\n"
    "def s_mean(x):\n"
    "    return s_sum(x) / len(x)\n\n"
    "def s_median(x):\n"
    "    s = sorted(x); n = len(s)\n"
    "    if n % 2 == 1:\n"
    "        return s[n // 2]\n"
    "    return (s[n // 2 - 1] + s[n // 2]) / 2.0\n\n"
    "def s_range(x):\n"
    "    return max(x) - min(x)\n\n"
    "def s_var_biased(x):\n"
    "    m = s_mean(x)\n"
    "    return sum((v - m) ** 2 for v in x) / len(x)\n\n"
    "def s_var_unbiased(x):\n"
    "    m = s_mean(x)\n"
    "    return sum((v - m) ** 2 for v in x) / (len(x) - 1)\n\n"
    "def s_initial_moment(x, k):\n"
    "    return sum(v ** k for v in x) / len(x)\n\n"
    "def s_central_moment(x, k):\n"
    "    m = s_mean(x)\n"
    "    return sum((v - m) ** k for v in x) / len(x)",
    "Листинг А.1 — Выборочные статистики",
)
add_blank()
add_listing(
    "def ecdf(sample):\n"
    "    s = sorted(sample)\n"
    "    n = len(s)\n"
    "    return s, [(i + 1) / n for i in range(n)]",
    "Листинг А.2 — Эмпирическая функция распределения",
)
add_blank()
add_listing(
    "import math\n"
    "def histogram(sample, bins=None):\n"
    "    n = len(sample)\n"
    "    if bins is None:\n"
    "        bins = max(1, int(round(1 + math.log2(n))))\n"
    "    a, b = min(sample), max(sample)\n"
    "    if a == b:\n"
    "        b = a + 1.0\n"
    "    edges = [a + (b - a) * i / bins for i in range(bins + 1)]\n"
    "    counts = [0] * bins\n"
    "    for v in sample:\n"
    "        idx = int((v - a) / (b - a) * bins)\n"
    "        if idx >= bins:\n"
    "            idx = bins - 1\n"
    "        counts[idx] += 1\n"
    "    dx = (b - a) / bins\n"
    "    density = [c / (n * dx) for c in counts]\n"
    "    return edges, counts, density",
    "Листинг А.3 — Гистограмма",
)
add_blank()
add_listing(
    "import math\n"
    "def loglaplace_pdf(x, mu, b):\n"
    "    if x <= 0:\n"
    "        return 0.0\n"
    "    return math.exp(-abs(math.log(x) - mu) / b) / (2 * b * x)\n\n"
    "def loglaplace_cdf(x, mu, b):\n"
    "    if x <= 0:\n"
    "        return 0.0\n"
    "    z = math.log(x) - mu\n"
    "    if z <= 0:\n"
    "        return 0.5 * math.exp(z / b)\n"
    "    return 1.0 - 0.5 * math.exp(-z / b)",
    "Листинг А.4 — Теоретическая плотность и функция распределения Log-Laplace",
)


# ---- Сохранение ----
out_path = ROOT / "Курсовая_ТВиМС_Вариант9.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}")
