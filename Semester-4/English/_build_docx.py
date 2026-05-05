"""Generate the final essay .docx with required formatting:
- Times New Roman 14, line spacing 1, margins 2 cm, first-line indent 1 cm
- Linkers highlighted yellow
- References list with italic titles
"""
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.section import WD_SECTION

ENGLISH_DIR = r"c:\Users\vyach\Study-Materials\Semester-4\English"
PLAGIARISM_CHECKS = [
    ("check_antiplagiat_ru.png", "Figure 1. Antiplagiat.ru — Originality 100%."),
    ("check_plagiat_ai.png", "Figure 2. Plagiat.ai — Originality 89%, AI-content low."),
    ("plagiarism_result.png", "Figure 3. PlagiarismDetector.net — Unique 100%, plagiarism not found."),
]

LINKERS = [
    "Furthermore", "As a result", "In addition", "Moreover", "In particular",
    "For instance", "Consequently", "On the other hand", "However",
    "Nevertheless", "Therefore", "To sum up", "In conclusion",
]

TITLE = "The Role of Blockchain Infrastructure in Shaping the Digital Society of the European Union"
GROUP = "Group 5151003/40001"
AUTHOR = "Totskiy Vjatseslav"
META = "Linkers: 13. Total: 576 words."

PARAGRAPHS = [
    'Blockchain technology is a decentralized digital ledger system that records transactions across a network of computers, ensuring transparency, security, and immutability of data. Over the past decade, it has evolved from the foundation of cryptocurrencies into a powerful tool for transforming public services and governance. The European Union has recognized this potential and placed blockchain at the core of its digital transformation strategy. <Furthermore>, the EU has invested significant resources into developing shared blockchain infrastructure to enhance cross-border cooperation among member states [1]. This essay argues that blockchain infrastructure plays a key role in building a secure, transparent, and efficient digital society in the European Union.',
    'The European Union has developed a comprehensive strategy for blockchain adoption. In 2018, all 27 EU member states, along with Norway and Liechtenstein, established the European Blockchain Partnership (EBP) to explore blockchain’s potential for the public sector [1]. <As a result>, the European Blockchain Services Infrastructure (EBSI) was launched. According to the European Commission, “EBSI is a peer-to-peer network of interconnected nodes running a blockchain-based services infrastructure” [1 – para. 1], designed to deliver cross-border public services across Europe. <In addition>, the EU introduced the Markets in Crypto-Assets Regulation (MiCA), which became fully applicable in December 2024. As ESMA states, “MiCA establishes uniform EU market rules for crypto-assets” [2 – para. 1], thereby providing legal certainty for the regulation and supervision of crypto-assets across all member states. <Moreover>, in May 2024, the European Commission established EUROPEUM-EDIC, a new legal entity tasked with governing EBSI and preparing the infrastructure for full-scale production [3]. These regulatory and institutional developments demonstrate the EU’s commitment to creating a structured and legally sound blockchain ecosystem.',
    'Blockchain technology offers numerous practical applications within EU digital services. <In particular>, one of the most significant use cases is digital identity verification through the concept of Self-Sovereign Identity (SSI), which empowers individuals to manage their own digital identities without relying on centralized authorities [4]. <For instance>, EBSI enables cross-border authentication of educational diplomas and professional certificates, allowing citizens to present verifiable credentials recognized across all EU member states. <Consequently>, this reduces bureaucratic procedures and eliminates the need for manual document verification. <On the other hand>, blockchain also enhances supply chain transparency by creating immutable records of product origins and movements, which is particularly important for food safety and pharmaceutical tracking within the European single market.',
    '<However>, despite these promising developments, several challenges hinder the widespread adoption of blockchain in the EU. Scalability remains a technical concern, as current blockchain networks struggle to process large volumes of transactions efficiently. Energy consumption is another issue, although EBSI addresses this by using a Proof of Authority consensus mechanism rather than energy-intensive Proof of Work [5]. <Nevertheless>, legal and regulatory fragmentation across EU member states presents additional obstacles, as different countries implement MiCA transitional measures at varying speeds [6]. <Therefore>, achieving full regulatory harmonization remains an ongoing process. Public trust and adoption barriers also persist, as many citizens and organizations lack sufficient understanding of blockchain technology to embrace it fully.',
    '<To sum up>, blockchain infrastructure has the potential to fundamentally reshape the digital society of the European Union by enabling secure identity management, cross-border document verification, and transparent supply chains. The EU has demonstrated strong commitment to this transformation through initiatives such as EBSI, the European Blockchain Partnership, and the MiCA regulation. <In conclusion>, while challenges related to scalability, regulatory fragmentation, and public adoption remain, the continued investment in blockchain infrastructure and balanced regulation will be essential for realizing the full potential of a decentralized digital Europe.',
]

REFERENCES = [
    ("European Commission. (2024). ", "European Blockchain Services Infrastructure", ". Shaping Europe’s Digital Future. Retrieved from https://digital-strategy.ec.europa.eu/en/policies/european-blockchain-services-infrastructure"),
    ("ESMA. (2025). ", "Markets in Crypto-Assets Regulation (MiCA)", ". Retrieved from https://www.esma.europa.eu/esmas-activities/digital-finance-and-innovation/markets-crypto-assets-regulation-mica"),
    ("European Commission. (2024). ", "About EUROPEUM-EDIC", ". Retrieved from https://ec.europa.eu/digital-building-blocks/sites/display/EBSI/About+us"),
    ("EQAR. (2025). ", "European Blockchain Service Infrastructure (EBSI)", ". Retrieved from https://www.eqar.eu/qa-results/synergies/european-blockchain-service-infrastructure-ebsi/"),
    ("MITA. (2025). ", "The European Blockchain Services Infrastructure (EBSI)", ". Retrieved from https://mita.gov.mt/2025/04/28/the-european-blockchain-services-infrastructure-ebsi/"),
    ("Hogan Lovells. (2025). ", "The EU’s Markets in Crypto-Assets MiCA Regulation — A Status Update", ". Retrieved from https://www.hoganlovells.com/en/publications/the-eus-markets-in-crypto-assets-mica-regulation-a-status-update"),
]


def set_run_font(run, size=14, bold=False, italic=False, highlight=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_body_paragraph(doc, text, *, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic_all=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = Cm(1)
    # Parse linkers in <...>
    parts = []
    buf = ""
    i = 0
    while i < len(text):
        if text[i] == "<":
            j = text.find(">", i)
            if j != -1:
                if buf:
                    parts.append(("text", buf))
                    buf = ""
                parts.append(("linker", text[i + 1 : j]))
                i = j + 1
                continue
        buf += text[i]
        i += 1
    if buf:
        parts.append(("text", buf))
    for kind, chunk in parts:
        run = p.add_run(chunk)
        set_run_font(run, italic=italic_all, highlight=(kind == "linker"))
    return p


def main():
    doc = Document()

    # Page margins 2 cm
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    # Title (centered, bold)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(TITLE)
    set_run_font(run, size=14, bold=True)

    # Group
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(GROUP)
    set_run_font(run)

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(AUTHOR)
    set_run_font(run)

    # Meta (italic)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(META)
    set_run_font(run, italic=True)

    # Body
    for para in PARAGRAPHS:
        add_body_paragraph(doc, para)

    # References heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("References")
    set_run_font(run, bold=True)

    # References items
    for idx, (prefix, italic_part, suffix) in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.left_indent = Cm(0.75)
        pf.first_line_indent = Cm(-0.75)
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        run = p.add_run(f"{idx}. {prefix}")
        set_run_font(run)
        run = p.add_run(italic_part)
        set_run_font(run, italic=True)
        run = p.add_run(suffix)
        set_run_font(run)

    # ---- Plagiarism check screenshots ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("Plagiarism Check Results")
    set_run_font(run, size=14, bold=True)

    for filename, caption in PLAGIARISM_CHECKS:
        path = os.path.join(ENGLISH_DIR, filename)
        if not os.path.exists(path):
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(path, width=Cm(16))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cap.paragraph_format.space_after = Pt(12)
        run = cap.add_run(caption)
        set_run_font(run, size=11, italic=True)

    out = r"c:\Users\vyach\Study-Materials\Semester-4\English\Essay_Blockchain_EU_Final.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
