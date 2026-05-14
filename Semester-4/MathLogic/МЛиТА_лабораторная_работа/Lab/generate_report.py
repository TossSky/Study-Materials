# -*- coding: utf-8 -*-
"""Генерация отчёта по лабораторной работе по МЛиТА.
Тема: «Составление программы для машины Тьюринга».
Вариант 6: f(x,y,z,t) = yzt + xyt + xyz.

Запуск:  python generate_report.py
Выход:   Отчёт_МЛиТА_Тоцкий_В_вар6.docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from itertools import product

BASE = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ---------- Глобальные стили ----------
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


# ---------- Помощники ----------
def add_centered(text, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.bold = bold
    r.font.size = Pt(size); r.font.name = 'Times New Roman'
    return p


def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(14); r.font.name = 'Times New Roman'
    return p


def add_subheading(text):
    add_blank()
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.name = 'Times New Roman'; r.font.size = Pt(14)
    return p


def add_body(text):
    return doc.add_paragraph(text)


def add_image(filename, width_cm=16, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run()
    r.add_picture(os.path.join(BASE, filename), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.first_line_indent = Cm(0)
        cr = cp.add_run(caption); cr.font.name = 'Times New Roman'; cr.font.size = Pt(14)


def add_table_label(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.font.size = Pt(14); r.font.name = 'Times New Roman'


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(h); r.font.name = 'Times New Roman'
        r.font.size = Pt(12); r.bold = True
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]; c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(v)); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    if col_widths:
        for row in t.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)


# ---------- Вычисление функции варианта 6 ----------
def f(x, y, z, t):
    return (y*z*t + x*y*t + x*y*z) % 2


# ========== ТИТУЛЬНАЯ СТРАНИЦА ==========
add_centered('Министерство науки и высшего образования Российской Федерации')
add_centered('Санкт-Петербургский политехнический университет Петра Великого')
add_blank()
add_centered('Институт компьютерных наук и кибербезопасности')
add_centered('Высшая школа кибербезопасности')
for _ in range(5):
    add_blank()
add_centered('ЛАБОРАТОРНАЯ РАБОТА', bold=True, size=16)
add_centered('«Составление программы для машины Тьюринга»')
add_blank()
add_centered('по дисциплине')
add_centered('«Математическая логика и теория алгоритмов»')
add_blank()
add_centered('Вариант 6')
for _ in range(5):
    add_blank()

for label, val in [('Выполнил',                          ''),
                   ('студент гр. 5151003/40001',        '\t\t\t\t\t\t\tТоцкий В. С.'),
                   ('',                                  ''),
                   ('Преподаватель',                     '\t\t\t\t\t\t\t\t\tПлатонов В. В.')]:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(label + val); r.font.size = Pt(14); r.font.name = 'Times New Roman'

for _ in range(4):
    add_blank()
add_centered('Санкт-Петербург')
add_centered('2026')
doc.add_page_break()


# ========== 1. ЦЕЛЬ И ЗАДАЧИ ==========
add_heading('1. Цель и задачи работы')
add_blank()
add_body(
    'Цель работы — закрепить теоретические сведения о машине Тьюринга '
    'как абстрактной модели алгоритма и приобрести практические навыки '
    'разработки программ для машины Тьюринга в одном из специализированных '
    'симуляторов (JFLAP).'
)
add_body(
    'Задача варианта 6: построить машину Тьюринга, вычисляющую значение '
    'двоичной булевой функции четырёх переменных, заданной многочленом '
    'Жегалкина:'
)
add_centered('f(x, y, z, t) = yzt + xyt + xyz   (mod 2).', bold=True)
add_body(
    'Исходные данные размещены на ленте в виде ␣xyzt␣ (символ ␣ обозначает '
    'пустую ячейку), где x, y, z, t ∈ {0, 1}; считывающая головка изначально '
    'установлена над первым символом исходных данных — над x. По завершении '
    'работы машины на ленте должен остаться единственный символ результата '
    '(0 или 1), все исходные символы должны быть стёрты, головка остановлена '
    'над результатом.'
)
doc.add_page_break()


# ========== 2. ХОД РАБОТЫ ==========
add_heading('2. Ход работы')

# 2.1 Таблица истинности
add_subheading('2.1. Таблица истинности функции варианта 6')
add_body(
    'Перед построением машины Тьюринга вычислены значения функции '
    'f(x, y, z, t) = yzt + xyt + xyz на всех 16 наборах входных переменных. '
    'Сложение в многочлене Жегалкина выполняется по модулю 2 (XOR). '
    'Результаты сведены в таблицу 1.'
)
add_blank()
add_table_label('Таблица 1 — Таблица истинности функции f(x, y, z, t) варианта 6')
truth_rows = []
for i, (x, y, z, t) in enumerate(product([0, 1], repeat=4)):
    truth_rows.append([
        i,
        f'{x}{y}{z}{t}',
        y*z*t,
        x*y*t,
        x*y*z,
        f(x, y, z, t),
    ])
add_table(
    ['№ набора', 'xyzt', 'yzt', 'xyt', 'xyz', 'f'],
    truth_rows,
    col_widths=[2.0, 2.0, 1.6, 1.6, 1.6, 1.6],
)
add_blank()
ones = [r[1] for r in truth_rows if r[5] == 1]
add_body(
    'Функция принимает значение 1 на четырёх наборах: ' +
    ', '.join(ones) + '. На остальных двенадцати наборах f = 0.'
)

# 2.2 Идея алгоритма
add_subheading('2.2. Идея алгоритма')
add_body(
    'Для четырёх входных битов существует ровно 16 различных наборов, что '
    'позволяет реализовать машину Тьюринга по схеме «дерева состояний»: '
    'на каждом шаге головка читает очередной символ исходных данных и '
    'переходит в состояние, кодирующее уже прочитанный префикс. Такой '
    'подход исключает необходимость в дополнительной разметке ленты '
    'служебными символами и в многократных проходах по входным данным.'
)
add_body(
    'Алгоритм состоит из двух фаз. На фазе чтения головка проходит четыре '
    'позиции исходных данных слева направо, не изменяя содержимое ленты, '
    'и накапливает информацию о значениях x, y, z, t в номере состояния. '
    'После чтения четвёртого символа (t) значение функции f(x, y, z, t) '
    'считается известным; машина стирает символ t и переходит в одно из '
    'двух «стирающих» состояний — erase0 или erase1, — несущих в своём '
    'имени уже вычисленный результат. На фазе очистки головка движется '
    'влево, последовательно стирая z, y, x; достигнув пустой ячейки слева '
    'от исходного x, машина записывает на её место символ результата '
    '(0 или 1) и переходит в финальное состояние qF без перемещения. '
    'На ленте остаётся единственный символ — значение функции — а головка '
    'стоит над ним, что в точности соответствует требованию задания.'
)

# 2.3 Состав состояний
add_subheading('2.3. Состав состояний')
add_body(
    'Машина Тьюринга включает 18 состояний (таблица 2). Имена состояний '
    'отражают их назначение: q0 — начальное; qx0/qx1 — после чтения x; '
    'qxyAB — после чтения x, y; qxyzABC — после чтения x, y, z; '
    'erase0/erase1 — стирающие проходы, несущие результат 0 или 1; '
    'qF — финальное состояние.'
)
add_blank()
add_table_label('Таблица 2 — Состав состояний машины Тьюринга')
add_table(
    ['Имя состояния', 'Назначение'],
    [
        ['q0',                   'Начальное состояние; головка над символом x'],
        ['qx0, qx1',             'Прочитан x; ожидается чтение y'],
        ['qxy00 … qxy11',        'Прочитаны x, y; ожидается чтение z (4 состояния)'],
        ['qxyz000 … qxyz111',    'Прочитаны x, y, z; ожидается чтение t (8 состояний)'],
        ['erase0',               'Стирание исходных данных; результат f = 0'],
        ['erase1',               'Стирание исходных данных; результат f = 1'],
        ['qF',                   'Финальное состояние; головка над результатом'],
    ],
    col_widths=[5.0, 11.5],
)

# 2.4 Внешний и ленточный алфавиты
add_subheading('2.4. Алфавиты и формальное определение машины')
add_body(
    'Внешний (входной) алфавит: Σ = {0, 1}. Ленточный алфавит: Γ = {0, 1, ␣}, '
    'где ␣ — пустая ячейка ленты. Множество состояний Q = {q0, qx0, qx1, '
    'qxy00, qxy01, qxy10, qxy11, qxyz000, …, qxyz111, erase0, erase1, qF} — '
    'всего 18 состояний. Начальное состояние q₀ = q0; единственное '
    'заключительное состояние qF. Множество направлений движения головки '
    '{L, R, S}.'
)

# 2.5 Таблица переходов
add_subheading('2.5. Таблица переходов')
add_body(
    'Поведение машины описывается программой переходов δ: Q × Γ → Q × Γ × {L, R, S}. '
    'Полный набор переходов представлен в таблице 3. В колонке «Запись» '
    'символ ␣ означает запись пустой ячейки (стирание); прочерк в колонке '
    '«Движение» при S означает, что головка остаётся на месте.'
)
add_blank()
add_table_label('Таблица 3 — Таблица переходов машины Тьюринга')

# Сборка таблицы переходов программно (соответствует turing_machine.jff)
def tr(src, sym, dst, write, move):
    return [src, sym, dst, write, move]

trans = []

# q0
trans += [tr('q0', '0', 'qx0', '0', 'R'),
          tr('q0', '1', 'qx1', '1', 'R')]
# qx*
for x in (0, 1):
    for y in (0, 1):
        trans += [tr(f'qx{x}', str(y), f'qxy{x}{y}', str(y), 'R')]
# qxy*
for x in (0, 1):
    for y in (0, 1):
        for z in (0, 1):
            trans += [tr(f'qxy{x}{y}', str(z), f'qxyz{x}{y}{z}', str(z), 'R')]
# qxyz* — читаем t, вычисляем f, стираем t, идём влево
for x in (0, 1):
    for y in (0, 1):
        for z in (0, 1):
            for t in (0, 1):
                r = f(x, y, z, t)
                trans += [tr(f'qxyz{x}{y}{z}', str(t),
                             f'erase{r}', '␣', 'L')]
# erase*: стираем 0/1 и идём влево; на ␣ записываем результат и стоп
for r in (0, 1):
    trans += [tr(f'erase{r}', '0', f'erase{r}', '␣', 'L'),
              tr(f'erase{r}', '1', f'erase{r}', '␣', 'L'),
              tr(f'erase{r}', '␣', 'qF', str(r), 'S')]

add_table(
    ['Состояние', 'Чтение', 'Новое состояние', 'Запись', 'Движение'],
    trans,
    col_widths=[3.5, 2.0, 3.7, 2.0, 2.6],
)

# 2.6 Реализация
add_subheading('2.6. Реализация в симуляторе JFLAP')
add_body(
    'Программа машины Тьюринга реализована в симуляторе JFLAP версии 7.1 '
    '(тип автомата — Turing Machine, одна лента). Файл проекта — '
    'turing_machine.jff. Граф автомата построен в соответствии с таблицей '
    'переходов: для каждой пары (состояние, прочитанный символ) добавлен '
    'переход с подписью вида «прочитанный_символ ; записываемый_символ , '
    'направление». Стирающие переходы помечены пустым символом ␣ в полях '
    'чтения или записи. Начальное состояние — q0 (отмечено стрелкой '
    '«Initial»), заключительное — qF (двойной обводкой, метка «Final»). '
    'Общий вид графа автомата представлен на рисунке 1.'
)
add_blank()
add_image('jflap_graph.png', width_cm=16,
          caption='Рисунок 1 — Граф машины Тьюринга в симуляторе JFLAP')
doc.add_page_break()

# 2.7 Тестирование
add_subheading('2.7. Тестирование')
add_body(
    'Корректность реализации проверена в режиме «Multiple Run (Transducer)» '
    'на всех шестнадцати возможных наборах входных данных (файл lab_test.txt '
    'из материалов задания). Результаты прогона представлены на рисунке 2.'
)
add_blank()
add_image('jflap_results.png', width_cm=14,
          caption='Рисунок 2 — Результаты прогона машины Тьюринга '
                  'на 16 входных наборах')
add_blank()
add_body(
    'Все 16 запусков завершились в заключительном состоянии qF (колонка '
    'Result — Accept). Значения в колонке Output полностью совпадают со '
    'значениями функции f(x, y, z, t) из таблицы истинности (таблица 1): '
    'на наборах 0111, 1101, 1110, 1111 машина возвращает 1; на остальных '
    'двенадцати наборах — 0. Лента после останова содержит единственный '
    'символ — значение функции; головка остановлена над ним. '
    'Это подтверждает корректность построенной машины Тьюринга.'
)
doc.add_page_break()


# ========== 3. ВЫВОДЫ ==========
add_heading('3. Выводы')
add_blank()

for i, s in enumerate([
    'Построена и реализована в симуляторе JFLAP машина Тьюринга, '
    'вычисляющая значение функции f(x, y, z, t) = yzt + xyt + xyz по '
    'четырём входным двоичным символам. Корректность машины подтверждена '
    'прогоном на полном множестве входных наборов (16 из 16) — результаты '
    'совпадают с независимо построенной таблицей истинности.',

    'Применённая «древовидная» схема состояний (1 + 2 + 4 + 8 = 15 '
    'состояний фазы чтения плюс 2 стирающих и 1 финальное, итого 18) '
    'обеспечивает однопроходное чтение исходных данных без вспомогательной '
    'разметки ленты. При фиксированной длине входа (4 символа) такая схема '
    'оказывается компактнее по числу переходов, чем универсальные '
    'конструкции, требующие хранения промежуточных значений на ленте.',

    'Вычислительная часть алгоритма (определение значения f) полностью '
    'выполняется за счёт ветвления управляющей логики — на этапе компиляции '
    'таблицы переходов; в сами правила перехода значения многочлена '
    'Жегалкина «зашиты» в выбор целевого стирающего состояния (erase0 либо '
    'erase1). Это иллюстрирует тезис о том, что машина Тьюринга может '
    'реализовать произвольную булеву функцию ограниченного числа аргументов '
    'без использования арифметики на ленте, перенося всю комбинаторику в '
    'граф состояний.',

    'Фаза очистки реализована минимально достаточным числом состояний (два): '
    'результат вычисления переносится в имя стирающего состояния, '
    'а конкретный символ ответа записывается единственным переходом по '
    'пустой ячейке. Это сохраняет инвариант задания «на ленте остаётся '
    'только результат» и обеспечивает остановку головки точно над символом '
    'ответа.',
], 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(f'{i}. '); r.font.name = 'Times New Roman'; r.font.size = Pt(14)
    r = p.add_run(s); r.font.name = 'Times New Roman'; r.font.size = Pt(14)


# ---------- Сохранение ----------
out = os.path.join(BASE, 'Отчёт_МЛиТА_Тоцкий_В_вар6.docx')
doc.save(out)
print(f'Saved: {out}')
