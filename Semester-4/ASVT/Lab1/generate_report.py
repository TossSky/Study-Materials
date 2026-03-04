# -*- coding: utf-8 -*-
"""Генерация отчёта по ЛР1, вариант 6а."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- Настройка стилей ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Поля страницы
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def add_centered_text(text, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_empty_line():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_heading_text(text):
    """Заголовок раздела (жирный, по центру)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p


def add_body_text(text):
    """Абзац основного текста"""
    p = doc.add_paragraph(text)
    return p


def add_subheading(text):
    """Подзаголовок (жирный, с отступом)"""
    add_empty_line()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p


def add_bold_and_text(bold_part, text_part, indent=True):
    """Абзац с жирным началом и обычным продолжением."""
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(bold_part)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run = p.add_run(text_part)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p


def add_list_item(text):
    """Элемент списка с тире."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'— {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p


def add_listing(code_text, caption):
    """Листинг кода в рамке (таблица из одной ячейки)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(caption)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ''
    code_lines = code_text.split('\n')
    for i, line in enumerate(code_lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)


def add_table_label(text):
    """Подпись таблицы — центр, не жирный."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'


def add_table_fmt(headers, rows):
    """Таблица: заголовки жирные 12pt, данные 12pt, всё по центру."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    return table


def add_formula(text):
    """Формула — по центру."""
    add_empty_line()
    add_centered_text(text)
    add_empty_line()


# ===== ТИТУЛЬНАЯ СТРАНИЦА =====
add_centered_text('Министерство науки и высшего образования Российской Федерации')
add_centered_text('Санкт-Петербургский политехнический университет Петра Великого')
add_empty_line()
add_centered_text('Институт компьютерных наук и кибербезопасности')
add_centered_text('Высшая школа кибербезопасности')

for _ in range(4):
    add_empty_line()

add_centered_text('ЛАБОРАТОРНАЯ РАБОТА №1', bold=True, size=16)
add_centered_text('«Основы работы в среде Atmel Studio.\nАрхитектура и система команд микроконтроллера ATmega32»')
add_empty_line()
add_centered_text('по дисциплине')
add_centered_text('«Аппаратные средства вычислительной техники»')

for _ in range(4):
    add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Выполнил')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('студент гр. 5151003/40001\t\t\t\t\t\t\tТоцкий В.')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Преподаватель\t\t\t\t\t\t\t\t\tМакаров А.С.')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

for _ in range(3):
    add_empty_line()

add_centered_text('Санкт-Петербург – 2026')

doc.add_page_break()

# ===== РАЗДЕЛ 1. СХЕМА УСТАНОВКИ =====
add_heading_text('1. Схема установки')
add_empty_line()

add_body_text(
    'Для выполнения лабораторной работы используется отладочная плата EasyAVR v7 '
    'с микроконтроллером ATmega32. Плата подключается к персональному компьютеру '
    'через USB-порт программатора mikroProg.'
)

add_body_text('Задействованные узлы отладочной платы:')

for c in [
    'микроконтроллер ATmega32 — центральный процессор, выполняющий программу;',
    'программатор mikroProg — для загрузки прошивки (.hex-файла) в МК через USB 2.0;',
    'кварцевый резонатор 8 МГц — задаёт тактовую частоту микроконтроллера;',
    'порты ввода-вывода PORTA–PORTD — настроены на вывод, подключены к светодиодам;',
    '32 светодиодных индикатора (по 8 на каждый порт) — визуализация состояния портов.',
]:
    add_list_item(c)

add_body_text(
    'Программа AVRFlash используется для записи .hex-файла в микроконтроллер. '
    'Все 4 порта ввода-вывода (PORTA, PORTB, PORTC, PORTD) настроены на выход '
    '(DDRx = 0xFF), к ним подключены светодиоды отладочной платы для визуального '
    'наблюдения эффекта «бегущего огня».'
)

doc.add_page_break()

# ===== РАЗДЕЛ 2. БЛОК-СХЕМА АЛГОРИТМА =====
add_heading_text('2. Блок-схема алгоритма работы программы')
add_empty_line()

add_subheading('2.1. Основная программа')

add_body_text('Алгоритм основной программы состоит из следующих этапов:')

for i, s in enumerate([
    'Инициализация 32-разрядного числа R3:R2:R1:R0 = 0x00000001 (единица в младшем бите).',
    'Настройка всех портов ввода-вывода на выход (DDRA–DDRD = 0xFF).',
    'Установка указателя стека SP на конец ОЗУ (RAMEND = 0x085F).',
    'Основной цикл (loop): циклический сдвиг 32-битного числа вправо на 1 бит, '
    'вывод в порты PORTA–PORTD, вызов задержки delay, переход к началу цикла.',
], 1):
    add_body_text(f'{i}. {s}')

add_subheading('2.2. Подпрограмма задержки (delay)')

add_body_text(
    'Подпрограмма delay реализует программную задержку '
    'с использованием двух вложенных циклов:'
)

for i, s in enumerate([
    'Загрузка начальных значений: R30 = y (внешний счётчик), R29 = x (внутренний счётчик).',
    'Внутренний цикл: INC R29 → NOP → BRNE delay_sub. Повторяется до переполнения R29 (R29 = 0).',
    'После выхода из внутреннего цикла: NOP → INC R30 → BRNE delay_sub. '
    'Если R30 ≠ 0, повторяется весь цикл.',
    'После выхода из внешнего цикла: NOP (дополнительный) → RET (возврат из подпрограммы).',
], 1):
    add_body_text(f'{i}. {s}')

doc.add_page_break()

# ===== РАЗДЕЛ 3. КОММЕНТИРОВАННЫЙ ЛИСТИНГ =====
add_heading_text('3. Комментированный листинг программы')
add_empty_line()

add_body_text(
    'В листинге 1 представлен полный текст программы на языке ассемблера '
    'для микроконтроллера ATmega32 с комментариями.'
)
add_empty_line()

code_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'lab1_var6a.asm')
with open(code_path, 'r', encoding='utf-8') as f:
    code = f.read()

add_listing(code, 'Листинг 1 — Исходный код программы (вариант 6а)')

doc.add_page_break()

# ===== РАЗДЕЛ 4. РАСЧЁТ ЗАДЕРЖКИ =====
add_heading_text('4. Расчёт задержки и определение констант x и y')
add_empty_line()

add_subheading('4.1. Анализ цикла задержки (вариант 6а)')

add_body_text(
    'Фрагмент кода задержки варианта 6а содержит два вложенных цикла. '
    'Внутренний цикл состоит из команд INC R29, NOP и BRNE. '
    'Внешний цикл — из NOP, INC R30 и BRNE.'
)

add_body_text('В таблице 1 приведено количество тактов для каждой команды.')

add_table_label('Таблица 1 — Количество тактов команд блока задержки')

add_table_fmt(
    ['Команда', 'Такты'],
    [
        ['LDI', '1'],
        ['INC', '1'],
        ['NOP', '1'],
        ['BRNE (переход выполнен)', '2'],
        ['BRNE (переход не выполнен)', '1'],
        ['RET', '4'],
    ])

add_subheading('4.2. Вывод формулы')

add_body_text(
    'Внутренний цикл (INC + NOP + BRNE) выполняется до тех пор, пока R29 не '
    'станет равен 0. Каждая итерация с переходом занимает 1 + 1 + 2 = 4 такта, '
    'последняя итерация (без перехода) — 1 + 1 + 1 = 3 такта.'
)

add_body_text(
    'Первый проход внутреннего цикла: R29 инкрементируется от x до 0 (через 255), '
    'что составляет (256 − x) итераций. Число тактов первого прохода:'
)

add_formula('N_вн.1 = (256 − x − 1) · 4 + 3 = 4 · (256 − x) − 1')

add_body_text(
    'Последующие проходы (R29 начинает с 0): 256 итераций, '
    'N_вн.x = 255 · 4 + 3 = 1023 такта.'
)

add_body_text(
    'Внешний цикл: R30 инкрементируется от y до 0, всего M = (256 − y) итераций. '
    'После каждого выхода из внутреннего цикла выполняются NOP + INC R30 + BRNE: '
    '4 такта (с переходом) или 3 такта (последняя итерация).'
)

add_body_text(
    'С учётом n дополнительных команд NOP перед RET, общая формула имеет вид:'
)

add_formula('N = 4 · (256 − x) + 1035 + (254 − y) · 1027 + n')

add_body_text(
    'где x — начальное значение R29, y — начальное значение R30, '
    'n — количество дополнительных NOP (от 0 до 4).'
)

add_subheading('4.3. Расчёт для задержки 10 мс')

add_body_text('Требуемое количество тактов при тактовой частоте 8 МГц:')

add_formula('N = t · f = 0,01 · 8 000 000 = 80 000 тактов')

add_body_text('Подставляя в формулу:')
add_body_text('80000 = 4 · (256 − x) + 1035 + (254 − y) · 1027 + n')
add_body_text('78965 − n = 4 · (256 − x) + (254 − y) · 1027')

add_body_text('Для n = 1 и (254 − y) = 76, то есть y = 178:')
add_body_text('4 · (256 − x) = 78964 − 76 · 1027 = 78964 − 78052 = 912')
add_body_text('256 − x = 228, откуда x = 28.')

add_body_text(
    'Проверка: N = 4·228 + 1035 + 76·1027 + 1 = 912 + 1035 + 78052 + 1 = 80000.'
)

add_body_text('Результаты расчёта представлены в таблице 2.')

add_table_label('Таблица 2 — Результаты расчёта параметров задержки')

add_table_fmt(
    ['Параметр', 'Значение'],
    [
        ['x (R29)', '28'],
        ['y (R30)', '178'],
        ['n (доп. NOP)', '1'],
        ['N (такты)', '80 000'],
        ['t (время)', '10 мс'],
    ])

doc.add_page_break()

# ===== РАЗДЕЛ 5. АНАЛИЗ .LSS-ФАЙЛА =====
add_heading_text('5. Анализ .lss-файла')
add_empty_line()

add_subheading('5.1. Адреса меток программы')

add_body_text(
    'В таблице 3 приведены адреса меток программы, полученные из .lss-файла '
    '(в словах, так как память программ ATmega32 адресуется 16-битными словами).'
)

add_table_label('Таблица 3 — Адреса меток программы')

add_table_fmt(
    ['Метка', 'Адрес (слова)', 'Адрес (байты)'],
    [
        ['(вектор сброса)', '0x0000', '0x0000'],
        ['delay', '0x0002', '0x0004'],
        ['delay_sub', '0x0004', '0x0008'],
        ['reset', '0x000C', '0x0018'],
        ['loop', '0x001B', '0x0036'],
        ['task9_label', '0x0028', '0x0050'],
    ])

add_body_text(
    'Команда JMP reset занимает 2 машинных слова (32 бита), поэтому подпрограмма '
    'delay начинается с адреса 0x0002. Метка reset расположена по адресу 0x000C '
    '(после 10 команд подпрограммы delay). Программа занимает 44 машинных слова (88 байт).'
)

add_subheading('5.2. Используемые форматы команд')

add_body_text('В таблице 4 перечислены форматы команд, задействованных в программе.')

add_table_label('Таблица 4 — Форматы команд программы')

add_table_fmt(
    ['Команда', 'Размер', 'Тип адресации', 'Операнды'],
    [
        ['JMP k', '32 бит', 'Прямая', 'k — 16-бит адрес'],
        ['LDI Rd, K', '16 бит', 'Непосредственная', 'd∈[16;31], K — 8 бит'],
        ['INC Rd', '16 бит', 'Регистровая', 'd — 5-бит номер'],
        ['NOP', '16 бит', 'Безоперандная', '—'],
        ['BRNE k', '16 бит', 'Относительная', 'k — 7-бит смещение'],
        ['RET', '16 бит', 'Безоперандная', '—'],
        ['MOV Rd, Rr', '16 бит', 'Двухрегистровая', 'd, r — 5 бит'],
        ['CLR Rd', '16 бит', 'Регистровая', 'EOR Rd, Rd'],
        ['SER Rd', '16 бит', 'Непосредственная', 'LDI Rd, 0xFF'],
        ['OUT A, Rr', '16 бит', 'Порт + регистр', 'A — 6 бит, r — 5 бит'],
        ['BST Rr, b', '16 бит', 'Регистр + бит', 'r — 5, b — 3 бит'],
        ['BLD Rd, b', '16 бит', 'Регистр + бит', 'd — 5, b — 3 бит'],
        ['LSR Rd', '16 бит', 'Регистровая', 'd — 5-бит номер'],
        ['ROR Rd', '16 бит', 'Регистровая', 'd — 5-бит номер'],
        ['CALL k', '32 бит', 'Прямая', 'k — 16-бит адрес'],
        ['RJMP k', '16 бит', 'Относительная', 'k — 12-бит смещение'],
        ['SET (BSET s)', '16 бит', 'Безоперандная', 's — 3-бит номер флага'],
        ['BREQ k', '16 бит', 'Относительная', 'k — 7-бит смещение'],
    ])

doc.add_page_break()

# ===== РАЗДЕЛ 6. АНАЛИЗ .HEX-ФАЙЛА =====
add_heading_text('6. Анализ .hex-файла')
add_empty_line()

add_body_text(
    'Сформированный Atmel Studio .hex-файл имеет формат Intel HEX. '
    'Структура каждой строки (записи) представлена в таблице 5.'
)

add_table_label('Таблица 5 — Структура записи Intel HEX')

add_table_fmt(
    ['Поле', 'Размер', 'Описание'],
    [
        [':', '1 символ', 'Маркер начала записи (RECORD MARK)'],
        ['RECLEN', '1 байт', 'Количество байт данных в записи'],
        ['LOAD OFFSET', '2 байта', 'Начальный адрес загрузки данных'],
        ['RECTYP', '1 байт', 'Тип: 00 — данные, 01 — конец файла, 02 — расш. адрес сегмента'],
        ['DATA', 'n байт', 'Байты данных программы'],
        ['CHKSUM', '1 байт', 'Контрольная сумма (доп. до двух)'],
    ])

add_body_text(
    'Контрольная сумма вычисляется как дополнение до двух суммы всех байт записи '
    '(RECLEN + старший и младший байты LOAD OFFSET + RECTYP + все байты DATA).'
)

add_body_text(
    'Программа содержит 44 машинных слова (88 байт). '
    '.hex-файл состоит из 8 записей: 1 запись расширенного адреса сегмента '
    '(RECTYP = 02), 6 записей данных (RECTYP = 00, из них 5 по 16 байт '
    'и 1 на 8 байт) и 1 запись конца файла (RECTYP = 01).'
)

add_body_text('Содержимое .hex-файла:')

hex_content = (
    ':020000020000FC\n'
    ':100000000C940C00E2EBDCE1D3950000E9F7000072\n'
    ':10001000E395D1F70000089541E0042E4427142E03\n'
    ':10002000242E342E4FEF4ABB47BB44BB41BB48E0B4\n'
    ':100030004EBF4FE54DBF00FA3694279417940794AE\n'
    ':1000400037F80BBA18BA25BA32BA0E940200F3CFB9\n'
    ':08005000000000006894E1F3D8\n'
    ':00000001FF'
)
add_listing(hex_content, 'Листинг 8 — Содержимое .hex-файла')

add_body_text(
    'Пример разбора первой записи данных: '
    'RECLEN = 0x10 (16 байт), LOAD OFFSET = 0x0000, RECTYP = 00 (данные), '
    'DATA = 0C940C00... (начиная с команды JMP reset = 0x940C 0x000C в формате '
    'little-endian), CHKSUM = 0x72.'
)

doc.add_page_break()

# ===== РАЗДЕЛ 7. ПОРЯДОК ВЫПОЛНЕНИЯ SBIS =====
add_heading_text('7. Порядок выполнения команды SBIS PORTA, 0')
add_empty_line()

add_subheading('7.1. Описание команды')

add_body_text(
    'SBIS PORTA, 0 — Skip if Bit in I/O Register is Set (пропуск следующей команды, '
    'если бит 0 регистра порта PORTA установлен в 1).'
)

add_body_text('Характеристики команды приведены в таблице 6.')

add_table_label('Таблица 6 — Характеристики команды SBIS PORTA, 0')

add_table_fmt(
    ['Характеристика', 'Значение'],
    [
        ['Мнемоника', 'SBIS'],
        ['Операнды', 'P = PORTA (I/O адрес 0x1B = 27), b = 0'],
        ['Формат КОП', '1001.1011.AAAA.Abbb'],
        ['Машинное слово', '1001.1011.1101.1000 = 0x9BD8'],
        ['Операция', 'if (I/O(P, b) = 1) then PC ← PC + 2 or 3'],
        ['Флаги', 'Не изменяет'],
        ['Такты', '1 / 2 / 3'],
    ])

add_subheading('7.2. Этапы выполнения команды в ЦП')

cpu_steps = [
    ('1. Указание адреса команды. ',
     'Из счётчика команд (PC) передаётся адрес текущей команды в память программ (ПЗУП).'),
    ('2. Увеличение PC. ',
     'Счётчик команд увеличивается на 1: PC ← PC + 1.'),
    ('3. Выборка команды. ',
     '16-битное машинное слово (0x9BD8) считывается из ПЗУП и записывается в регистр команды.'),
    ('4. Декодирование. ',
     'Устройство управления распознаёт формат 1001.1011.AAAA.Abbb как SBIS. '
     'Извлекается адрес порта A = 11011₂ = 27 (PORTA) и номер бита b = 000₂ = 0.'),
    ('5. Чтение порта ввода-вывода. ',
     'Из регистра PORTA (I/O адрес 0x1B, адрес в памяти 0x3B) считывается значение. '
     'Из 8-битного значения выделяется бит 0.'),
    ('6. Проверка условия. ',
     'Устройство управления проверяет: бит 0 PORTA == 1?'),
    ('7. Обновление PC. ',
     'Если бит = 0: PC не изменяется (выполняется следующая команда, 1 такт). '
     'Если бит = 1: PC ← PC + 1 (пропуск 16-бит команды, 2 такта) '
     'или PC ← PC + 2 (пропуск 32-бит команды, 3 такта).'),
    ('8. Переход к следующей команде. ',
     'Цикл повторяется с пункта 1.'),
]
for bold_part, text_part in cpu_steps:
    add_bold_and_text(bold_part, text_part)

add_subheading('7.3. Задействованные узлы МК')

add_body_text(
    'В таблице 7 перечислены узлы МК, задействованные при выполнении команды SBIS.'
)

add_table_label('Таблица 7 — Задействованные узлы МК')

add_table_fmt(
    ['Узел МК', 'Назначение'],
    [
        ['Счётчик команд (PC)', 'Адрес текущей/следующей команды'],
        ['ПЗУП (Flash)', 'Хранение машинного кода'],
        ['Регистр команды', 'Хранение текущей инструкции'],
        ['Устройство управления', 'Декодирование, управляющие сигналы'],
        ['Порт PORTA', 'Источник проверяемого бита'],
    ])

doc.add_page_break()

# ===== РАЗДЕЛ 8. ДЕКОДИРОВАНИЕ МАШИННЫХ СЛОВ =====
add_heading_text('8. Декодирование машинных слов')
add_empty_line()

add_subheading('8.1. Машинное слово 0x9468')

add_body_text('Двоичное представление: 1001.0100.0110.1000.')

add_body_text(
    'Сравнение с форматом: 1001.0100.0sss.1000 — '
    'команда BSET s (установка флага в SREG).'
)

add_body_text(
    'Извлечение поля s: биты [6:4] = 110₂ = 6. '
    'Бит 6 регистра SREG — флаг T (Transfer bit, бит пересылки).'
)

add_body_text(
    'Результат: SET (BSET 6) — установка флага пересылки T. Операция: T ← 1.'
)

add_subheading('8.2. Машинное слово 0xF3E1')

add_body_text('Двоичное представление: 1111.0011.1110.0001.')

add_body_text(
    'Сравнение с форматом: 1111.00kk.kkkk.ksss — '
    'команда BRBS s, k (переход при установленном флаге).'
)

add_body_text(
    'Извлечение полей: s = биты [2:0] = 001₂ = 1 (флаг Z — Zero). '
    'k = биты [9:3] = 1111100₂. В 7-битном дополнительном коде: k = −4.'
)

add_body_text(
    'Результат: BREQ −4 (BRBS 1, −4) — переход если Z = 1, смещение −4. '
    'Операция: if (Z = 1) then PC ← PC + (−4) + 1 = PC − 3.'
)

add_subheading('8.3. Сводная таблица')

add_body_text('В таблице 8 представлены результаты декодирования.')

add_table_label('Таблица 8 — Результаты декодирования машинных слов')

add_table_fmt(
    ['Маш. слово', 'Двоичный код', 'Мнемоника', 'Описание'],
    [
        ['0x9468', '1001.0100.0110.1000', 'SET (BSET 6)', 'Установка флага T'],
        ['0xF3E1', '1111.0011.1110.0001', 'BREQ −4', 'Переход если Z=1, k=−4'],
    ])

add_body_text(
    'Для проверки команды размещены в конце программы (метка task9_label). '
    'Перед SET добавлены два NOP для обеспечения смещения BREQ = −4.'
)

add_body_text(
    'Сравнение с .lss-файлом: ассемблер формирует машинное слово SET = 0x9468, '
    'что совпадает с заданием. BREQ task9_label формирует 0xF3E1, '
    'что также совпадает с заданием. Оба декодированных слова верны.'
)

doc.add_page_break()

# ===== РАЗДЕЛ 9. АЛГОРИТМ ВЫПОЛНЕНИЯ КОМАНД =====
add_heading_text('9. Алгоритм выполнения команд ассемблера')
add_empty_line()

add_body_text(
    'Ниже описаны все команды ассемблера, задействованные в программе, '
    'с указанием формата, операции и количества тактов.'
)

cmds = [
    ('JMP k', ' — прямой безусловный переход. Формат: 32 бита. Операция: PC ← k. '
     'Этапы: выборка 2 слов из ПЗУП → декодирование → запись адреса k в PC. 3 такта.'),
    ('LDI Rd, K', ' — загрузка константы в регистр. Формат: 16 бит, d∈[16;31]. '
     'Операция: Rd ← K. Этапы: выборка → декодирование → извлечение K из команды → '
     'запись K в Rd. 1 такт.'),
    ('INC Rd', ' — инкремент регистра. Операция: Rd ← Rd + 1. Этапы: выборка → '
     'декодирование → чтение Rd → передача в АЛУ → прибавление 1 → запись результата → '
     'обновление флагов Z, N, V, S. 1 такт.'),
    ('NOP', ' — пустая команда. Машинное слово: 0x0000. '
     'Этапы: выборка → декодирование → инкремент PC. 1 такт.'),
    ('BRNE k', ' — условный переход при Z = 0 (алиас BRBC 1, k). Этапы: выборка → '
     'декодирование → чтение Z из SREG → если Z = 0: вычисление нового PC = PC + k + 1. '
     '1 такт (без перехода) / 2 такта (с переходом).'),
    ('RET', ' — возврат из подпрограммы. Операция: PC ← Stack. Этапы: выборка → '
     'декодирование → чтение 2 байт из стека (ОЗУ) → запись в PC → инкремент SP на 2. '
     '4 такта.'),
    ('MOV Rd, Rr', ' — копирование регистра. Операция: Rd ← Rr. 1 такт.'),
    ('CLR Rd', ' — очистка регистра. Транслируется как EOR Rd, Rd. '
     'Операция: Rd ← 0x00. 1 такт.'),
    ('SER Rd', ' — установка всех битов. Транслируется как LDI Rd, 0xFF. '
     'Операция: Rd ← 0xFF. 1 такт.'),
    ('OUT A, Rr', ' — запись регистра в порт. Операция: I/O(A) ← Rr. '
     'Этапы: выборка → декодирование → чтение Rr → запись в порт ввода-вывода. 1 такт.'),
    ('BST Rr, b', ' — сохранение бита регистра во флаг T. '
     'Операция: T ← Rr(b). 1 такт.'),
    ('BLD Rd, b', ' — загрузка флага T в бит регистра. '
     'Операция: Rd(b) ← T. 1 такт.'),
    ('LSR Rd', ' — логический сдвиг вправо. '
     'Операция: C ← Rd(0), Rd(n) ← Rd(n+1), Rd(7) ← 0. 1 такт.'),
    ('ROR Rd', ' — циклический сдвиг вправо через перенос. '
     'Операция: C ← Rd(0), Rd(n) ← Rd(n+1), Rd(7) ← C_old. 1 такт.'),
    ('CALL k', ' — прямой вызов подпрограммы. Формат: 32 бита. '
     'Операция: Stack ← PC+2, PC ← k. 4 такта.'),
    ('RJMP k', ' — относительный безусловный переход. '
     'Операция: PC ← PC + k + 1. 2 такта.'),
    ('SET (BSET 6)', ' — установка флага T в регистре SREG. '
     'Формат: 1001.0100.0sss.1000, s = 6. Операция: T ← 1. 1 такт.'),
    ('BREQ k', ' — условный переход при Z = 1 (алиас BRBS 1, k). Этапы: выборка → '
     'декодирование → чтение Z из SREG → если Z = 1: вычисление нового PC = PC + k + 1. '
     '1 такт (без перехода) / 2 такта (с переходом).'),
]
for cmd, desc in cmds:
    add_bold_and_text(cmd, desc)

doc.add_page_break()

# ===== РАЗДЕЛ 10. ОТВЕТЫ НА КОНТРОЛЬНЫЕ ВОПРОСЫ =====
add_heading_text('10. Ответы на контрольные вопросы')
add_empty_line()

# --- Вопрос 1 ---
p = doc.add_paragraph()
run = p.add_run(
    '1) Основные узлы микроконтроллера ATmega32 и их назначение.'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_empty_line()

nodes = [
    ('АЛУ', ' — выполняет арифметические и логические операции над данными из РОН.'),
    ('Регистры общего назначения (32 шт.)',
     ' — быстрая память для операндов; R26–R31 образуют пары X, Y, Z для косвенной адресации.'),
    ('Счётчик команд (PC)', ' — хранит адрес текущей команды в ПЗУП.'),
    ('Регистр команды', ' — хранит текущую инструкцию из памяти программ.'),
    ('Регистр SREG', ' — содержит флаги результатов: C, Z, N, V, S, H, T, I.'),
    ('ПЗУП (Flash, 32 КБ)', ' — энергонезависимая память программ.'),
    ('ОЗУ (SRAM, 2 КБ)', ' — оперативная память данных и стек.'),
    ('EEPROM (1 КБ)', ' — энергонезависимая память данных.'),
    ('Порты PORTA–PORTD',
     ' — 32 линии ввода-вывода с регистрами DDRx, PORTx, PINx.'),
    ('Таймеры-счётчики',
     ' — два 8-битных (Timer0, Timer2) и один 16-битный (Timer1).'),
    ('Блок прерываний', ' — обработка внешних и внутренних событий.'),
    ('Интерфейсы USART, SPI, I2C', ' — последовательный обмен данными.'),
    ('АЦП (8 каналов, 10 бит)', ' — аналого-цифровое преобразование.'),
    ('Аналоговый компаратор', ' — сравнение аналоговых напряжений.'),
    ('Сторожевой таймер', ' — автоматический сброс МК при зависании программы.'),
]
for name, desc in nodes:
    add_bold_and_text('— ' + name, desc, indent=False)

add_empty_line()

# --- Вопрос 2 ---
p = doc.add_paragraph()
run = p.add_run('2) Признаки RISC-архитектуры в ATmega32.')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_empty_line()

add_body_text('Признаки RISC-архитектуры:')
for s in [
    'фиксированная длина команд (16 бит, некоторые 32 бита);',
    'большой регистровый файл (32 РОН);',
    'большинство команд выполняется за 1 такт;',
    'операции только над регистрами (Load/Store для памяти);',
    'гарвардская архитектура (раздельные шины программ и данных).',
]:
    add_list_item(s)

add_body_text(
    'Преимущества: высокая производительность (до 16 MIPS), предсказуемое время '
    'выполнения, простота конвейеризации, низкое энергопотребление.'
)

add_body_text(
    'Недостатки: больший объём кода по сравнению с CISC, ограничения на номера '
    'регистров и диапазоны констант из-за фиксированной длины команды.'
)

add_empty_line()

# --- Вопрос 3 ---
p = doc.add_paragraph()
run = p.add_run('3) Время выполнения команд CPSE и BRHS.')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_empty_line()

add_body_text(
    'Команда CPSE Rd, Rr (Compare, Skip if Equal): 1 такт если Rd ≠ Rr (пропуска нет); '
    '2 такта если Rd = Rr и следующая команда 16-битная; '
    '3 такта если Rd = Rr и следующая команда 32-битная.'
)

add_body_text('Пример (1 такт — без пропуска):')
add_listing(
    '    LDI  R16, 5       ; R16 = 5\n'
    '    LDI  R17, 10      ; R17 = 10\n'
    '    CPSE R16, R17     ; R16 != R17, пропуска нет (1 такт)\n'
    '    NOP               ; выполняется',
    'Листинг 2 — Пример CPSE без пропуска')

add_empty_line()
add_body_text('Пример (2 такта — пропуск 16-бит команды):')
add_listing(
    '    LDI  R16, 5       ; R16 = 5\n'
    '    LDI  R17, 5       ; R17 = 5\n'
    '    CPSE R16, R17     ; R16 = R17, пропуск NOP (2 такта)\n'
    '    NOP               ; пропускается\n'
    '    INC  R16          ; выполняется',
    'Листинг 3 — Пример CPSE с пропуском')

add_empty_line()

add_body_text(
    'Команда BRHS k (Branch if Half Carry Set): 1 такт если H = 0 (перехода нет); '
    '2 такта если H = 1 (переход выполняется).'
)

add_body_text('Пример (1 такт — без перехода):')
add_listing(
    '    CLH               ; H = 0\n'
    '    BRHS skip         ; H = 0, перехода нет (1 такт)\n'
    '    NOP               ; выполняется\n'
    'skip:',
    'Листинг 4 — Пример BRHS без перехода')

add_empty_line()
add_body_text('Пример (2 такта — с переходом):')
add_listing(
    '    LDI  R16, 0x0F    ; R16 = 0x0F\n'
    '    LDI  R17, 0x01    ; R17 = 0x01\n'
    '    ADD  R16, R17     ; 0x0F + 0x01 = 0x10, перенос из бита 3 -> H=1\n'
    '    BRHS skip         ; H = 1, переход выполняется (2 такта)\n'
    '    NOP               ; пропускается\n'
    'skip:',
    'Листинг 5 — Пример BRHS с переходом')

add_empty_line()

# --- Вопрос 4 ---
p = doc.add_paragraph()
run = p.add_run('4) Циклический сдвиг и сложение с константой.')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_empty_line()

add_body_text(
    'Циклический сдвиг вправо 8-разрядного числа (бит 0 переходит в бит 7) '
    'без использования флага T (для регистров R16–R31):'
)

add_listing(
    '    LSR  Rd           ; C <- Rd(0), Rd >>= 1, Rd(7) = 0\n'
    '    BRCC skip         ; если C = 0, бит 7 уже 0 - готово\n'
    '    ORI  Rd, 0x80     ; иначе установить бит 7\n'
    'skip:',
    'Листинг 6 — Циклический сдвиг вправо')

add_empty_line()

add_body_text(
    'Сложение 24-разрядного числа (R18:R17:R16) с 24-разрядной константой K '
    '(три команды). Поскольку команды ADDI не существует, используется '
    'вычитание отрицательной константы:'
)

add_listing(
    '    SUBI R16, LOW(-K)     ; R16 <- R16 + K_low  (вычитание -K)\n'
    '    SBCI R17, HIGH(-K)    ; R17 <- R17 + K_mid - C\n'
    '    SBCI R18, BYTE3(-K)   ; R18 <- R18 + K_high - C',
    'Листинг 7 — Сложение с 24-битной константой')

add_empty_line()

# --- Вопрос 5 ---
p = doc.add_paragraph()
run = p.add_run('5) Причины ограничений номеров регистров и констант.')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_empty_line()

add_body_text(
    'Все ограничения обусловлены фиксированной длиной машинного слова (16 бит) '
    'и необходимостью разместить в нём код операции и все операнды:'
)

for l in [
    'd∈[16;31] в LDI, ANDI, ORI, SUBI, SBCI, CPI — на номер регистра отведено '
    '4 бита (а не 5), что позволяет адресовать только 16 из 32 регистров;',
    'K∈[0;63] в ADIW/SBIW — константа кодируется 6 битами;',
    'dl∈{24,26,28,30} в ADIW/SBIW — на номер регистровой пары отведено всего '
    '2 бита (4 варианта);',
    'k∈[−64;63] в условных переходах (BRNE, BREQ и др.) — смещение кодируется '
    '7 битами в дополнительном коде;',
    'P∈[0;31] в CBI, SBI, SBIC, SBIS — адрес порта кодируется 5 битами, '
    'что ограничивает доступ первыми 32 портами.',
]:
    add_list_item(l)

add_body_text(
    'Таким образом, ограничения являются прямым следствием выбора RISC-архитектуры '
    'с фиксированной длиной команды: чем больше бит отводится на один операнд, '
    'тем меньше остаётся для других операндов и кода операции.'
)

# --- Сохранение ---
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    'Отчёт_ЛР1_Вариант6а.docx'
)
doc.save(output_path)
print(f'Отчёт сохранён: {output_path}')
