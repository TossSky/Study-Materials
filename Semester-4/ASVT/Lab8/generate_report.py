# -*- coding: utf-8 -*-
"""Генерация отчёта по ЛР8 (АСВТ, вариант 6а): прототип системы шифрования RSA
на ATmega32 + ПК-клиент через UART.
Запуск: python generate_report.py
Выход:  Отчёт_ЛР8_Вариант6а.docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ---------- Глобальные стили ----------
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


# ---------- Помощники (идентичны ЛР1) ----------
def add_centered(text, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.bold = bold
    r.font.size = Pt(size); r.font.name = 'Times New Roman'
    return p

def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    return p

def add_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(14); r.font.name = 'Times New Roman'
    return p

def add_subheading(text):
    add_blank()
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.name = 'Times New Roman'; r.font.size = Pt(14)
    return p

def add_body(text):
    return doc.add_paragraph(text)

def add_list_item(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    r = p.add_run(f'— {text}')
    r.font.name = 'Times New Roman'; r.font.size = Pt(14)
    return p

def add_listing(code, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(caption); r.font.size = Pt(14); r.font.name = 'Times New Roman'
    table = doc.add_table(rows=1, cols=1); table.style = 'Table Grid'
    cell = table.cell(0, 0); cell.text = ''
    for i, line in enumerate(code.split('\n')):
        pp = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        pp.paragraph_format.first_line_indent = Cm(0)
        pp.paragraph_format.line_spacing = 1.0
        pp.paragraph_format.space_after = Pt(0)
        pp.paragraph_format.space_before = Pt(0)
        pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = pp.add_run(line); rr.font.name = 'Courier New'; rr.font.size = Pt(9)

def add_table_label(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.font.size = Pt(14); r.font.name = 'Times New Roman'

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(h); r.font.name = 'Times New Roman'
        r.font.size = Pt(12); r.bold = True
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]; c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(v)); r.font.name = 'Times New Roman'; r.font.size = Pt(12)


# ========== ТИТУЛЬНАЯ СТРАНИЦА ==========
add_centered('Министерство науки и высшего образования Российской Федерации')
add_centered('Санкт-Петербургский политехнический университет Петра Великого')
add_blank()
add_centered('Институт компьютерных наук и кибербезопасности')
add_centered('Высшая школа кибербезопасности')
for _ in range(4): add_blank()
add_centered('ЛАБОРАТОРНАЯ РАБОТА №8', bold=True, size=16)
add_centered('«Взаимодействие микроконтроллера и ПЭВМ.\nПрототип системы шифрования»')
add_blank()
add_centered('по дисциплине')
add_centered('«Аппаратные средства вычислительной техники»')
add_blank()
add_centered('Вариант 6а')
for _ in range(4): add_blank()

for label, val in [('Выполнил', ''),
                   ('студент гр. 5151003/40001', '\t\t\t\t\t\t\tТоцкий В.'),
                   ('', ''),
                   ('Преподаватель', '\t\t\t\t\t\t\t\t\tМакаров А.С.')]:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(label + val); r.font.size = Pt(14); r.font.name = 'Times New Roman'

for _ in range(3): add_blank()
add_centered('Санкт-Петербург – 2026')
doc.add_page_break()


# ========== 1. ЦЕЛЬ РАБОТЫ ==========
add_heading('1. Цель работы')
add_blank()
add_body(
    'Получение практических навыков по организации двунаправленного '
    'взаимодействия между микроконтроллером ATmega32 и персональным компьютером. '
    'Знакомство с криптосистемой RSA. Разработка программного прототипа системы '
    'шифрования, состоящего из прошивки микроконтроллера и консольного приложения '
    'для ПК, взаимодействующих через UART-виртуальный COM-порт.'
)
add_body(
    'Согласно варианту 6, микроконтроллер должен хранить в EEPROM несколько троек '
    'ключей (n, e, d), поддерживать шифрование, расшифрование, подписание и '
    'проверку подписи; при расшифровании и подписании выполнять поиск секретного '
    'ключа d по открытой паре (n, e); вести статистику (суммарная длина '
    'зашифрованных/расшифрованных файлов, число подписаний и проверок подписи) '
    'и отображать её на 7-сегментном индикаторе в виде E.xxx, d.xxx, S.xxx, C.xxx '
    'с переключением режима по внешнему прерыванию INT0.'
)
add_body(
    'Параметры для проверки работоспособности: n = 39913, e = 453, d = 1221 '
    '(p = 167, q = 239, φ(n) = 39508; проверка: 453 · 1221 mod 39508 = 1).'
)
doc.add_page_break()


# ========== 2. СХЕМА УСТАНОВКИ ==========
add_heading('2. Схема установки')
add_blank()
add_body(
    'Для выполнения работы используется отладочная плата EasyAVR v7 '
    'с микроконтроллером ATmega32, работающим на частоте 8 МГц. '
    'Взаимодействие с ПК осуществляется через встроенный преобразователь '
    'USB-UART (виртуальный COM-порт), настроенный на 9600 бит/с, 8N1.'
)
add_body('Задействованные узлы отладочной платы приведены в таблице 1.')

add_table_label('Таблица 1 — Узлы отладочной платы')
add_table(
    ['Узел МК', 'Назначение'],
    [
        ['PD0 (RXD), PD1 (TXD)', 'Обмен с ПК по UART, 9600 8N1'],
        ['PD2 (INT0)', 'Кнопка переключения режима индикации (по падающему фронту)'],
        ['PORTC [7:0]', 'Сегменты a..g + dp 7-сегментного индикатора (общий катод)'],
        ['PORTA [3:0]', 'Выбор разряда индикатора (динамическая мультиплексирующая индикация)'],
        ['EEPROM (1 Кбайт)', 'Хранение массива ключей (до 4 троек) и блока статистики'],
    ])
add_blank()
add_body(
    'На стороне ПК используется консольное приложение, открывающее виртуальный '
    'COM-порт через WinAPI (CreateFile("\\\\\\\\.\\\\COMx")).'
)
doc.add_page_break()


# ========== 3. БЛОК-СХЕМЫ АЛГОРИТМОВ ==========
add_heading('3. Блок-схемы алгоритмов')
add_blank()

add_subheading('3.1. Основной цикл микроконтроллера')
add_body(
    'После сброса выполняется инициализация UART, 7-сегментного индикатора '
    'и внешнего прерывания INT0, затем разрешаются прерывания (команда sei). '
    'В основном цикле программа чередует два действия: неблокирующий опрос '
    'флага UCSRA.RXC (готовность приёма байта по UART) и мультиплексированный '
    'вывод одного разряда индикатора.'
)
for i, s in enumerate([
    'Проверка флага RXC: если в регистре UDR есть байт — вызывается функция '
    'dispatch(), которая блокирующе дочитывает оставшуюся часть команды '
    'и отправляет ответ.',
    'Загрузка блока статистики из EEPROM.',
    'Выбор отображаемого параметра (E/d/S/C) согласно значению disp_mode.',
    'Вывод одного разряда индикатора через функцию display_show(): для каждого '
    'разряда — установка сегментов на PORTC, выбор разряда на PORTA, задержка '
    '2 мс вызовом ассемблерной подпрограммы delay_2ms_asm().',
], 1):
    add_body(f'{i}. {s}')

add_subheading('3.2. Обработчик команды DECRYPT (как пример потоковой обработки)')
for i, s in enumerate([
    'Принять заголовок: n (2 байта), e (2 байта), len (2 байта, число 16-битных '
    'блоков шифртекста).',
    'Выполнить поиск записи (n, e) в EEPROM. Если не найдена — отправить '
    'на ПК байт статуса ST_NO_KEY (0xE1) и завершить обработку команды.',
    'Отправить байт статуса ST_OK (0x00).',
    'Цикл len раз: принять 2 байта шифртекста c, вычислить m = c^d mod n, '
    'отправить один байт m. Микроконтроллер не буферизует весь файл — '
    'обработка идёт побайтно, что соответствует требованию методички.',
    'Обновить счётчик dec_len в блоке статистики и записать блок в EEPROM.',
], 1):
    add_body(f'{i}. {s}')

add_subheading('3.3. Приложение для ПК')
for i, s in enumerate([
    'Открыть виртуальный COM-порт, настроить параметры (9600 8N1, тайм-ауты).',
    'Отобразить меню команд (add, del, list, enc, dec, sign, ver, stats, rstats).',
    'Для каждой введённой команды: сериализовать параметры в big-endian, '
    'записать в порт, прочитать байт статуса.',
    'Если OK — прочитать/записать потоковую часть (шифртекст или открытый текст, '
    'подпись, статистику) и вывести результат пользователю.',
    'В случае ошибки — вывести код статуса и вернуться к меню.',
], 1):
    add_body(f'{i}. {s}')

doc.add_page_break()


# ========== 4. ФОРМАТЫ ПЕРЕСЫЛАЕМЫХ ДАННЫХ ==========
add_heading('4. Форматы пересылаемых данных')
add_blank()
add_body(
    'Протокол взаимодействия ПК и МК построен поверх простого байтового потока '
    'UART. Все многобайтные поля передаются в формате big-endian. Каждая команда '
    'от ПК начинается с одного байта кода команды (CMD), ответ микроконтроллера — '
    'с одного байта статуса. В командах шифрования и расшифрования после '
    'заголовка следует потоковая часть переменной длины.'
)

add_table_label('Таблица 2 — Состав команд протокола')
add_table(
    ['CMD', 'Имя', 'От ПК к МК', 'От МК к ПК'],
    [
        ['0x01', 'ADD_KEY',     'n(2) e(2) d(2)',           'status + slot(1)'],
        ['0x02', 'DEL_KEY',     'slot(1)',                  'status'],
        ['0x03', 'LIST_KEYS',   '—',                         'status + count(1) + count·(n(2) e(2))'],
        ['0x04', 'ENCRYPT',     'n(2) e(2) len(2) + поток', 'status + поток len·c(2)'],
        ['0x05', 'DECRYPT',     'n(2) e(2) len(2) + поток', 'status + поток len байт'],
        ['0x06', 'SIGN',        'n(2) e(2) hash(2)',         'status + sig(2)'],
        ['0x07', 'VERIFY',      'n(2) e(2) hash(2) sig(2)', 'status + result(1)'],
        ['0x08', 'GET_STATS',   '—',                         'status + enc(4) dec(4) sign(2) check(2)'],
        ['0x09', 'RESET_STATS', '—',                         'status'],
    ])
add_blank()

add_table_label('Таблица 3 — Коды статуса')
add_table(
    ['Код', 'Константа', 'Смысл'],
    [
        ['0x00', 'ST_OK',          'Операция выполнена успешно'],
        ['0xE1', 'ST_NO_KEY',      'Секретный ключ d не найден по (n, e) в EEPROM'],
        ['0xE2', 'ST_BAD_PARAM',   'Некорректный параметр (например, hash ≥ n, slot вне диапазона)'],
        ['0xE3', 'ST_FULL',        'В EEPROM нет свободных слотов для нового ключа'],
        ['0xEF', 'ST_INVALID_CMD', 'Неизвестный код команды'],
    ])
add_blank()

add_subheading('Ключевая особенность — потоковая обработка')
add_body(
    'Для команд ENCRYPT/DECRYPT микроконтроллер не буферизует весь файл. После '
    'приёма заголовка он читает один байт открытого текста (одно 16-битное слово '
    'шифртекста), выполняет модульное возведение и немедленно отправляет '
    'результат, после чего переходит к следующему блоку. Это удовлетворяет '
    'требованию методички: «недопустима предварительная передача всего файла на '
    'МК с последующей обработкой и возвратом полностью сформированного '
    'зашифрованного/расшифрованного файла».'
)
doc.add_page_break()


# ========== 5. ЛИСТИНГИ ==========
add_heading('5. Комментированные листинги')
add_blank()

add_subheading('5.1. Математика RSA (rsa.c)')
add_body(
    'Модульное возведение в степень выполняется бинарным методом (right-to-left), '
    'что даёт O(log exp) умножений. Поскольку n² ≈ 1,6·10⁹ не помещается в 16 бит, '
    'промежуточные произведения хранятся в uint32_t.'
)
add_blank()

BASE = os.path.dirname(os.path.abspath(__file__))

def read(rel):
    with open(os.path.join(BASE, rel), 'r', encoding='utf-8') as f:
        return f.read()

add_listing(read('rsa.c'), 'Листинг 1 — Модульная арифметика (rsa.c)')
add_blank()

add_subheading('5.2. Прошивка микроконтроллера (mcu.c)')
add_body(
    'Модуль объединяет все функции прошивки: драйвер UART (9600 8N1, UBRR = 51 '
    'при F_CPU = 8 МГц), работу с EEPROM через avr/eeprom.h, динамическую '
    'индикацию 7-сегментного дисплея, обработчик внешнего прерывания INT0 '
    '(переключение режима отображения), диспетчер команд протокола и основной '
    'цикл main().'
)
add_blank()
add_listing(read('mcu.c'), 'Листинг 2 — Прошивка ATmega32 (mcu.c)')
doc.add_page_break()

add_subheading('5.3. Подпрограмма задержки на ассемблере (delay_asm.S)')
add_body(
    'Динамическая индикация 7-сегментного дисплея требует задержки между '
    'переключениями разрядов — в нашей реализации 2 мс на разряд. Для получения '
    'ассемблерного листинга в прошивке эта задержка выполнена отдельной '
    'подпрограммой на языке ассемблера (файл delay_asm.S), линкуемой вместе '
    'с Си-кодом и вызываемой из функции display_show() по внешнему имени '
    'delay_2ms_asm().'
)
add_body(
    'Структура подпрограммы аналогична ЛР1 вариант 6а: два вложенных цикла '
    '(внутренний — INC R29, NOP, BRNE; внешний — NOP, INC R30, BRNE). '
    'Требуемое число тактов при F_CPU = 8 МГц — 16 000; подобранные значения '
    'x = 110 (R29), y = 240 (R30), n = 3 (дополнительных NOP) дают формулу '
    'N = 4·(256 − 110) + 1035 + (254 − 240)·1027 + 3 = 16 000.'
)
add_blank()
add_listing(read('delay_asm.S'), 'Листинг 3 — Подпрограмма задержки (delay_asm.S)')
doc.add_page_break()

add_subheading('5.4. Приложение для ПК (pc_client.c)')
add_body(
    'Консольный клиент для Windows, реализует открытие виртуального COM-порта '
    '(CreateFile("\\\\\\\\.\\\\COMx"), настройка DCB и COMMTIMEOUTS), сериализацию '
    'команд протокола и пользовательское меню. Программа хэширует файл '
    'алгоритмом CRC-16-CCITT перед подписанием и проверкой подписи, редуцируя '
    'результат по модулю n.'
)
add_blank()
add_listing(read('pc_client.c'), 'Листинг 4 — Приложение для ПК (pc_client.c)')
doc.add_page_break()


# ========== 6. ВЫВОДЫ ==========
add_heading('6. Выводы')
add_blank()

for i, s in enumerate([
    'Разработан и протестирован программный прототип криптосистемы RSA, '
    'удовлетворяющий требованиям варианта 6: хранение нескольких троек ключей '
    'в EEPROM, поиск секретного ключа по открытой паре, побайтная (потоковая) '
    'обработка данных, ведение и отображение статистики на 7-сегментном '
    'индикаторе с переключением по внешнему прерыванию INT0.',

    'Выбранная архитектура с выделением всей платформо-независимой логики '
    '(модульная арифметика RSA, обработчик протокола) в отдельные модули через '
    'HAL-абстракцию оказалась полезной не только для переносимости, но и для '
    'тестирования: один и тот же код работает как в прошивке ATmega32, так и '
    'в настольном автотесте без единой строки дубляжа.',

    'При n = 39913 операция m^e mod n принципиально умещается в 16-битный '
    'регистр результата, однако требует 32-битного промежуточного произведения — '
    'именно поэтому в rsa_modexp используется uint32_t. На реальных размерах '
    'ключей RSA (1024 бит и более) такая схема неприменима: потребовалась бы '
    'библиотека длинной арифметики (например, GMP) и криптографически стойкий '
    'хэш (SHA-256 вместо CRC-16).',

    'Передача длины открытого текста в заголовке команды (как и указано в '
    'теоретическом описании протокола Диффи-Хеллмана в методичке) оказалась '
    'естественным способом решения проблемы границ сообщения в потоковом '
    'UART-протоколе без разделителей.',

    'Для защиты от атак повторного воспроизведения и перебора в реальной '
    'системе протокола, помимо цифровой подписи, потребовались бы метки времени '
    'или одноразовые номера (nonce) и аутентификация канала. В учебном прототипе '
    'эти механизмы сознательно не реализованы.',
], 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(f'{i}. '); r.font.name = 'Times New Roman'; r.font.size = Pt(14)
    r = p.add_run(s); r.font.name = 'Times New Roman'; r.font.size = Pt(14)


# ---------- Сохранение ----------
out = os.path.join(BASE, 'Отчёт_ЛР8_Вариант6а.docx')
doc.save(out)
print(f'Saved: {out}')
