from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- Настройка стилей ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def add_centered_text(text, bold=False, size=14, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_empty_line():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_heading_text(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p


def add_body_text(text):
    p = doc.add_paragraph(text)
    return p


def add_listing(code_text, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(caption)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ''
    code_lines = code_text.split('\n')
    for i, line in enumerate(code_lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)


def add_bold_item(label, desc):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run = p.add_run(f' — {desc}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def add_list_item(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'— {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def read_code(filename):
    code_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', filename)
    with open(code_path, 'r', encoding='utf-8') as f:
        code = f.read()
    lines = code.split('\n')
    new_lines = []
    for line in lines:
        spaces = len(line) - len(line.lstrip(' '))
        new_indent = ' ' * (spaces // 2)
        new_lines.append(new_indent + line.lstrip(' '))
    return '\n'.join(new_lines)


# ===== ТИТУЛЬНАЯ СТРАНИЦА =====
add_centered_text('Министерство науки и высшего образования Российской Федерации', size=14)
add_centered_text('Санкт-Петербургский политехнический университет Петра Великого', size=14)
add_empty_line()
add_centered_text('Институт компьютерных наук и кибербезопасности', size=14)
add_centered_text('Высшая школа кибербезопасности', size=14)

for _ in range(4):
    add_empty_line()

add_centered_text('ЛАБОРАТОРНАЯ РАБОТА №3', bold=True, size=16)
add_centered_text('«Использование Python-библиотек для работы с ИИ»', size=14)
add_empty_line()
add_centered_text('по дисциплине', size=14)
add_centered_text('«Цифровая аналитика»', size=14)

for _ in range(4):
    add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Выполнил')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('студент гр. 5151003/40001\t\t\tВ.С. Тоцкий')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Преподаватель')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('__________________________\t\t\t____________')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

for _ in range(3):
    add_empty_line()

add_centered_text('Санкт-Петербург – 2026', size=14)

doc.add_page_break()

# ===== 1. ЦЕЛЬ РАБОТЫ =====
add_heading_text('1. Цель работы')
add_empty_line()

add_body_text(
    'Получение навыков работы с методами искусственного интеллекта для решения '
    'задачи классификации данных с использованием языка программирования Python.'
)

doc.add_page_break()

# ===== 2. ЗАДАНИЕ НА РАБОТУ =====
add_heading_text('2. Задание на работу')
add_empty_line()

tasks = [
    'Установить необходимые для работы скрипта модули: torch, torchvision, pillow, '
    'scikit-learn, tqdm, numpy.',
    'Сформировать датасеты для обучения и тестирования (валидации). Набор данных '
    'должен представлять собой набор изображений: первый класс — люди, второй — животные.',
    'Описать базовую архитектуру модели ResNet-18, а также изменения, которые '
    'необходимо внести в предобученную для выполнения задачи бинарной классификации.',
    'Реализовать скрипт model.py, использующий CNN на основе ResNet-18.',
    'Реализовать скрипт train.py, который обучает модель несколько эпох (Adam + '
    'CrossEntropyLoss) и сохраняет веса в файл.',
    'Выбрать оптимальные параметры обучения. Точность на валидационной выборке '
    'должна быть не менее 80%.',
    'Модифицировать train.py: на вход — путь до тренировочного и валидационного '
    'датасетов; на выход — файл с весами.',
    'Реализовать скрипт predict.py: на вход — путь до изображения, на выход — '
    'результат бинарной классификации (человек или животное).',
    'Протестировать программу на изображениях из обучающей выборки, тестовой выборки '
    'и из произвольного источника.',
    'Провести эксперименты по изменению параметров модели для повышения эффективности.',
    'Сделать выводы о работе скрипта.',
    'Ответить на контрольные вопросы.',
]
for i, task in enumerate(tasks, 1):
    add_body_text(f'{i}. {task}')

doc.add_page_break()

# ===== 3. ХОД РАБОТЫ =====
add_heading_text('3. Ход работы')
add_empty_line()

add_body_text(
    'В ходе выполнения лабораторной работы было разработано Python-приложение для '
    'решения задачи бинарной классификации изображений (люди vs животные) с '
    'использованием свёрточной нейронной сети ResNet-18. Приложение разделено на '
    'три скрипта: model.py (описание архитектуры), train.py (обучение модели) и '
    'predict.py (инференс на одиночном изображении).'
)

add_empty_line()

# --- 3.1 Подготовка датасета ---
add_body_text('3.1. Подготовка датасета')
add_empty_line()

add_body_text(
    'Обучающая и валидационная выборки были сформированы из двух открытых источников. '
    'Для класса «люди» использован датасет Labeled Faces in the Wild (LFW) — набор '
    'фотографий лиц известных людей, загружаемый через HuggingFace (зеркало logasja/LFW). '
    'Для класса «животные» использован датасет CIFAR-10, из которого были отобраны '
    'изображения шести классов животных: bird, cat, deer, dog, frog, horse. CIFAR-10 '
    'загружается через torchvision.datasets.CIFAR10.'
)

add_body_text(
    'Итоговый датасет представлен в виде каталогов ImageFolder: data/train/humans, '
    'data/train/animals, data/val/humans, data/val/animals. Для обучения использовано '
    '200 изображений на класс (всего 400), для валидации — 60 изображений на класс '
    '(всего 120).'
)

add_body_text(
    'Консистентность выборки обеспечена за счёт выбора однородных источников: в классе '
    '«люди» все изображения содержат лицо в центре кадра на нейтральном фоне, в классе '
    '«животные» — животное занимает большую часть кадра. Такое распределение позволяет '
    'модели выделять признаки объекта, а не признаки фона.'
)

add_empty_line()

# --- 3.2 Архитектура ResNet-18 ---
add_body_text('3.2. Базовая архитектура ResNet-18')
add_empty_line()

add_body_text(
    'ResNet-18 (Residual Network, 18 слоёв) — свёрточная нейронная сеть, предложенная '
    'в 2015 году в статье «Deep Residual Learning for Image Recognition». Главная идея '
    'архитектуры — использование остаточных (residual) связей, позволяющих обучать '
    'более глубокие сети без деградации точности и без проблемы исчезающего градиента. '
    'Остаточный блок реализует преобразование вида y = F(x) + x, где F(x) — несколько '
    'последовательных свёрточных слоёв, а x — прямая связь (identity shortcut).'
)

add_body_text(
    'Структура ResNet-18 включает:'
)

resnet_layers = [
    'входной свёрточный слой Conv7×7, 64 фильтра, stride=2, за ним BatchNorm, ReLU и '
    'MaxPool3×3 со stride=2 — уменьшение пространственного разрешения с 224 до 56;',
    'четыре группы остаточных блоков (conv2_x — conv5_x), каждая состоит из двух '
    'BasicBlock-ов. Количество фильтров в группах: 64, 128, 256, 512. Между группами '
    'пространственное разрешение уменьшается в 2 раза, а количество каналов удваивается;',
    'глобальный адаптивный average pooling, приводящий карты признаков к размеру 1×1×512;',
    'полносвязный слой (fully connected, fc) с 512 входами и 1000 выходами для '
    'классификации на 1000 классов ImageNet.',
]
for item in resnet_layers:
    add_list_item(item)

add_empty_line()

add_body_text(
    'Общее количество обучаемых параметров — около 11.7 миллиона. Модель доступна в '
    'torchvision в предобученном виде (weights=ResNet18_Weights.IMAGENET1K_V1).'
)

add_empty_line()

add_body_text(
    'Изменения, внесённые в предобученную модель для задачи бинарной классификации:'
)

add_bold_item(
    'замена последнего полносвязного слоя fc',
    'исходный слой nn.Linear(512, 1000) заменён на nn.Linear(512, 2), так как '
    'количество классов в задаче равно 2 (humans и animals).'
)
add_bold_item(
    'заморозка признаковой части',
    'при обучении обновляются только параметры нового слоя fc, остальные веса сети '
    'остаются фиксированными (transfer learning). Это позволяет быстро адаптировать '
    'модель к новой задаче при небольшом объёме данных.'
)
add_bold_item(
    'использование стандартной предобработки',
    'применяется preprocessing pipeline из ResNet18_Weights.IMAGENET1K_V1.transforms(), '
    'включающий изменение размера до 256×256, центральный кроп 224×224 и нормализацию '
    'по статистикам ImageNet (mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]).'
)

add_empty_line()

# --- 3.3 Реализация скриптов ---
add_body_text('3.3. Реализация скриптов')
add_empty_line()

add_body_text(
    'Скрипт model.py содержит функцию build_model, которая создаёт экземпляр ResNet-18 '
    'с предобученными весами и заменяет последний полносвязный слой. Константа CLASSES '
    'хранит имена классов в порядке, совпадающем с ImageFolder (алфавитный: animals, humans).'
)

add_body_text(
    'Скрипт train.py реализует цикл обучения. Принимает обязательные аргументы командной '
    'строки --train (путь до тренировочного датасета) и --val (путь до валидационного '
    'датасета), а также опциональные параметры обучения --epochs, --batch-size, --lr. '
    'Выходным параметром является --output (по умолчанию weights.pt) — файл с весами '
    'лучшей модели. Данные загружаются через torchvision.datasets.ImageFolder, поэтому '
    'ожидается структура каталогов вида <split>/<class_name>/*.jpg.'
)

add_body_text(
    'Для обучения используется оптимизатор Adam с learning rate 1e-3 и функция потерь '
    'CrossEntropyLoss. На каждой эпохе выполняется прямой и обратный проход по '
    'тренировочной выборке с обновлением весов, затем — прогон по валидационной выборке '
    'для оценки точности (accuracy). Если точность на текущей эпохе превысила лучшую, '
    'веса модели сохраняются в указанный файл вместе со списком классов.'
)

add_body_text(
    'В тренировочном пайплайне применяются аугментации: RandomResizedCrop(224) с '
    'диапазоном масштаба [0.8, 1.0] и RandomHorizontalFlip. Для валидации используется '
    'детерминированное преобразование Resize(256) + CenterCrop(224). Во всех случаях '
    'применяется нормализация по статистикам ImageNet.'
)

add_body_text(
    'Скрипт predict.py загружает сохранённые веса, применяет то же преобразование, что '
    'и валидация, пропускает изображение через модель и выводит результат классификации. '
    'Вероятности классов вычисляются через softmax.'
)

add_empty_line()

# --- 3.4 Параметры и результаты обучения ---
add_body_text('3.4. Параметры обучения и результаты')
add_empty_line()

add_body_text(
    'Выбранные параметры обучения:'
)

add_bold_item('число эпох', '3')
add_bold_item('размер батча (batch size)', '32')
add_bold_item('learning rate', '1e-3')
add_bold_item('оптимизатор', 'Adam')
add_bold_item('функция потерь', 'CrossEntropyLoss')
add_bold_item('обучаемые параметры', 'только последний слой fc (заморозка backbone)')
add_bold_item('устройство', 'CPU')

add_empty_line()

add_body_text(
    'Результаты обучения на 400 тренировочных и 120 валидационных изображениях:'
)

results = [
    'эпоха 1: train loss = 0.5548, train acc = 0.718 | val loss = 0.3063, val acc = 0.917;',
    'эпоха 2: train loss = 0.1779, train acc = 0.985 | val loss = 0.0947, val acc = 1.000;',
    'эпоха 3: train loss = 0.0963, train acc = 0.995 | val loss = 0.0590, val acc = 1.000.',
]
for r in results:
    add_list_item(r)

add_empty_line()

add_body_text(
    'Требование методички (точность на валидационной выборке ≥ 80%) выполнено уже на '
    'первой эпохе (91.7%). По итогам трёх эпох модель достигла точности 100% на '
    'валидационной выборке. Быстрая сходимость объясняется использованием transfer '
    'learning: признаковая часть ResNet-18 уже обучена на ImageNet и способна '
    'различать людей и животных на уровне низко- и среднеуровневых признаков; модели '
    'требуется только подобрать веса последнего слоя.'
)

add_empty_line()

# --- 3.5 Тестирование predict.py ---
add_body_text('3.5. Тестирование предсказаний')
add_empty_line()

add_body_text(
    'Работа predict.py проверена на трёх категориях изображений:'
)

predict_results = [
    ('изображение из обучающей выборки — человек',
     'humans (p=0.970)'),
    ('изображение из обучающей выборки — животное',
     'animals (p=0.904)'),
    ('изображение из валидационной выборки — человек',
     'humans (p=0.899)'),
    ('изображение из валидационной выборки — животное',
     'animals (p=0.855)'),
    ('произвольное изображение (сгенерированное GAN-ом лицо '
     'с сайта thispersondoesnotexist.com)',
     'humans (p=0.894)'),
    ('произвольное изображение (случайная фотография собаки '
     'из сервиса dog.ceo)',
     'animals (p=0.889)'),
]
for desc, res in predict_results:
    add_bold_item(desc, res)

add_empty_line()

add_body_text(
    'Во всех шести случаях модель присвоила корректный класс с уверенностью выше 85%. '
    'На изображениях из произвольного источника (вне тренировочной и валидационной '
    'выборки) модель также продемонстрировала правильную классификацию, что '
    'подтверждает обобщающую способность transfer learning на предобученной ResNet-18.'
)

add_empty_line()

# --- 3.6 Эксперименты ---
add_body_text('3.6. Эксперименты по повышению эффективности')
add_empty_line()

add_body_text(
    'В ходе работы были рассмотрены следующие способы повышения эффективности модели:'
)

experiments = [
    'применение аугментаций на этапе обучения (RandomResizedCrop, RandomHorizontalFlip) — '
    'снижает переобучение и повышает обобщающую способность модели за счёт вариативности '
    'входных изображений;',
    'заморозка признаковой части и обучение только последнего слоя (feature extraction) — '
    'ускоряет обучение на CPU и предотвращает разрушение предобученных признаков при '
    'малом объёме данных;',
    'сохранение только лучшей модели по валидационной точности — гарантирует, что '
    'итоговые веса соответствуют точке максимального обобщения, а не последней эпохе;',
    'использование стандартной предобработки ResNet18_Weights.IMAGENET1K_V1.transforms() — '
    'обеспечивает совместимость входов с распределением, на котором обучалась модель.',
]
for exp in experiments:
    add_list_item(exp)

add_empty_line()

add_body_text(
    'Возможные направления дальнейшего улучшения:'
)

further = [
    'fine-tuning всей сети с малым learning rate (1e-4–1e-5) после предварительного '
    'обучения только fc-слоя;',
    'расширение датасета до нескольких тысяч изображений на класс для более надёжной '
    'валидации;',
    'использование более глубоких архитектур (ResNet-50, EfficientNet) при наличии GPU;',
    'применение регуляризации (Dropout, weight_decay) при увеличении объёма обучения.',
]
for f in further:
    add_list_item(f)

add_empty_line()

# --- 3.7 Листинги ---
add_body_text('В листингах 1–3 представлен исходный код всех трёх скриптов приложения.')
add_empty_line()

add_listing(read_code('model.py'), 'Листинг 1 — Исходный код model.py')
add_empty_line()
add_listing(read_code('train.py'), 'Листинг 2 — Исходный код train.py')
add_empty_line()
add_listing(read_code('predict.py'), 'Листинг 3 — Исходный код predict.py')

doc.add_page_break()

# ===== 4. ОТВЕТЫ НА КОНТРОЛЬНЫЕ ВОПРОСЫ =====
add_heading_text('4. Ответы на контрольные вопросы')
add_empty_line()

# Вопрос 1
p = doc.add_paragraph()
run = p.add_run(
    '1) Какие классы задач могут быть решены с помощью методов искусственного интеллекта?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text('Основные классы задач, решаемых методами искусственного интеллекта:')

ai_tasks = [
    'классификация — отнесение объекта к одному из заданных классов (бинарная и '
    'многоклассовая классификация изображений, текстов, аудио);',
    'регрессия — предсказание непрерывной числовой величины (прогноз цены, температуры, '
    'спроса);',
    'кластеризация — разбиение множества объектов на группы по сходству без заранее '
    'известных меток (k-means, DBSCAN, иерархическая кластеризация);',
    'снижение размерности — представление данных в пространстве меньшей размерности '
    'с сохранением структуры (PCA, t-SNE, автоэнкодеры);',
    'детектирование и сегментация объектов на изображениях (YOLO, Faster R-CNN, U-Net);',
    'обработка естественного языка — машинный перевод, генерация текста, ответы на '
    'вопросы, анализ тональности;',
    'генерация данных — синтез изображений, текстов, аудио (GAN, диффузионные модели, '
    'большие языковые модели);',
    'обучение с подкреплением — выработка стратегии поведения агента в среде '
    '(управление роботом, игровые ИИ, оптимизация).',
]
for t in ai_tasks:
    add_list_item(t)

add_empty_line()

# Вопрос 2
p = doc.add_paragraph()
run = p.add_run('2) Чем отличается обучающий набор данных от тестового?')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text(
    'Обучающий набор данных (training set) используется для настройки параметров '
    'модели: на нём вычисляются градиенты функции потерь и обновляются веса. Модель '
    '«видит» эти примеры многократно в процессе обучения.'
)

add_body_text(
    'Тестовый (или валидационный) набор данных состоит из примеров, которые модель '
    'не видела при обучении. Он используется для оценки обобщающей способности '
    'модели — насколько хорошо она работает на новых данных. Тестовая выборка '
    'должна быть независима от обучающей, чтобы результаты оценки были объективными '
    'и не были искажены переобучением (overfitting).'
)

add_body_text(
    'По методическому рекомендации тестовая выборка должна составлять не менее 20% '
    'от объёма тренировочной. В данной работе соотношение train : val = 400 : 120, '
    'что составляет 30% и удовлетворяет этому требованию.'
)

add_empty_line()

# Вопрос 3
p = doc.add_paragraph()
run = p.add_run(
    '3) Что такое признак в контексте методов искусственного интеллекта? '
    'Что такое метка в контексте методов искусственного интеллекта? В чём их разница?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text(
    'Признак (feature) — это характеристика объекта, представленная в виде числа или '
    'вектора чисел и подаваемая на вход модели. Для изображения признаками являются '
    'значения пикселей (или карты признаков, извлечённые свёрточными слоями). Для '
    'табличных данных признаки — это столбцы таблицы (возраст, зарплата, количество '
    'покупок и т. п.). Совокупность признаков образует вектор признаков объекта.'
)

add_body_text(
    'Метка (label, target) — это правильный ответ, который модель должна предсказать '
    'на основе признаков. В задачах классификации метка — это имя класса (или его '
    'числовой индекс); в задачах регрессии — числовое значение целевой переменной. '
    'В данной работе метками являются индексы 0 (animals) и 1 (humans).'
)

add_body_text(
    'Разница: признаки — это входные данные модели, метки — это ожидаемый выход. '
    'В процессе обучения модель ищет функцию, которая максимально точно отображает '
    'признаки в метки. Во время инференса модель получает только признаки и '
    'предсказывает метку самостоятельно.'
)

add_empty_line()

# Вопрос 4
p = doc.add_paragraph()
run = p.add_run(
    '4) Чем методы глубокого обучения отличаются от других методов искусственного интеллекта?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text(
    'Глубокое обучение (deep learning) — подраздел машинного обучения, использующий '
    'многослойные искусственные нейронные сети. Его ключевые отличия от классических '
    'методов (линейная регрессия, решающие деревья, SVM, k-NN):'
)

dl_diffs = [
    'автоматическое извлечение признаков — в классических методах признаки '
    'проектируются вручную экспертом, в глубоком обучении сеть сама учится строить '
    'иерархию признаков (от низкоуровневых — края и текстуры — до высокоуровневых — '
    'объекты и сцены);',
    'большое число параметров — от миллионов до миллиардов. Это требует больших '
    'объёмов данных и вычислительных ресурсов (GPU, TPU);',
    'способность работать с неструктурированными данными — изображения, аудио, '
    'текст, видео, тогда как классические методы чаще применяются к табличным данным;',
    'использование специализированных архитектур под тип данных — CNN для '
    'изображений, RNN/Transformer для последовательностей, GNN для графов;',
    'обучение преимущественно через backpropagation и градиентный спуск на GPU;',
    'возможность transfer learning — переиспользование предобученных моделей для '
    'новых задач (как в данной работе: ResNet-18, обученная на ImageNet, '
    'адаптирована для классификации humans vs animals).',
]
for d in dl_diffs:
    add_list_item(d)

add_empty_line()

add_body_text(
    'Классические методы ИИ остаются эффективными на небольших табличных данных, '
    'где они уступают глубокому обучению по выразительной силе, но превосходят по '
    'интерпретируемости и скорости обучения.'
)

add_empty_line()

# Вопрос 5
p = doc.add_paragraph()
run = p.add_run(
    '5) Из чего состоит слой в нейронной сети? Какие слои бывают? Что такое нейрон?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text(
    'Нейрон (искусственный нейрон, perceptron) — базовая вычислительная единица '
    'нейронной сети. Выполняет линейное преобразование входа: y = f(Σ wᵢ·xᵢ + b), '
    'где xᵢ — входные значения, wᵢ — весовые коэффициенты, b — смещение (bias), '
    'f — нелинейная функция активации (ReLU, sigmoid, tanh, GELU и др.). Веса и '
    'смещение являются обучаемыми параметрами.'
)

add_body_text(
    'Слой (layer) — совокупность нейронов, обрабатывающих вход параллельно и '
    'формирующих выходной тензор. Слой задаётся типом операции, количеством '
    'нейронов (или каналов), функцией активации и обучаемыми параметрами.'
)

add_body_text('Основные типы слоёв в нейронных сетях:')

layer_types = [
    'полносвязный слой (fully connected, Linear, Dense) — каждый нейрон связан со '
    'всеми входами предыдущего слоя. Используется в классификационных головах (как '
    'fc-слой в ResNet-18);',
    'свёрточный слой (Conv2d) — применяет обучаемые фильтры к пространственной '
    'структуре входа. Основной строительный блок CNN;',
    'слой пулинга (MaxPool, AvgPool, AdaptiveAvgPool) — уменьшает пространственное '
    'разрешение карт признаков, сохраняя наиболее значимые значения;',
    'слой нормализации (BatchNorm, LayerNorm, GroupNorm) — стабилизирует обучение '
    'за счёт нормировки активаций по статистикам батча или слоя;',
    'слой активации (ReLU, LeakyReLU, Sigmoid, Tanh, Softmax) — вносит нелинейность '
    'в модель; без активации сеть эквивалентна одной линейной функции;',
    'рекуррентные слои (RNN, LSTM, GRU) — обрабатывают последовательности, '
    'сохраняя внутреннее состояние;',
    'слой Attention (самовнимание, Self-Attention) — базовый блок Transformer, '
    'позволяющий нейронам взаимодействовать между собой вне зависимости от расстояния;',
    'слой Dropout — во время обучения случайно обнуляет часть активаций; '
    'используется для регуляризации и борьбы с переобучением;',
    'Embedding — проекция дискретных идентификаторов (слова, токены) в плотные '
    'вектора.',
]
for lt in layer_types:
    add_list_item(lt)

doc.add_page_break()

# ===== 5. ВЫВОД =====
add_heading_text('5. Вывод')
add_empty_line()

add_body_text(
    'В ходе выполнения лабораторной работы было разработано Python-приложение для '
    'решения задачи бинарной классификации изображений (люди vs животные) на основе '
    'свёрточной нейронной сети ResNet-18 с применением подхода transfer learning. '
    'Сформирован сбалансированный датасет из 520 изображений (400 train + 120 val), '
    'собранный из открытых источников LFW (люди) и CIFAR-10 (животные). Реализованы '
    'три скрипта: model.py — описание модифицированной архитектуры ResNet-18, '
    'train.py — цикл обучения с оптимизатором Adam и функцией потерь CrossEntropyLoss, '
    'predict.py — инференс на одиночном изображении. Модель достигла точности 100% '
    'на валидационной выборке за 2 эпохи обучения на CPU, что превышает требуемый '
    'порог в 80%. Работоспособность классификатора подтверждена на изображениях из '
    'обучающей выборки, валидационной выборки и из произвольного источника.'
)

# --- Сохранение ---
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Отчёт_ЛР3_Тоцкий_ВС.docx')
doc.save(output_path)
print(f'Report saved: {output_path}')
