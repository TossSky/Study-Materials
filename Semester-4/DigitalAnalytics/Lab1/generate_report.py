from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# --- Настройка стилей ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Поля страницы
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def add_centered_text(text, bold=False, size=14, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_empty_line():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_heading_text(text):
    """Добавить заголовок раздела (жирный, по центру)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p


def add_body_text(text):
    """Добавить абзац основного текста"""
    p = doc.add_paragraph(text)
    return p


def add_listing(code_text, caption):
    """Добавить листинг кода в рамке (таблица из одной ячейки)"""
    # Название листинга
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(caption)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    # Таблица-рамка для кода
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ''
    # Разбиваем код на строки и добавляем каждую как отдельный параграф
    code_lines = code_text.split('\n')
    for i, line in enumerate(code_lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)


# ===== ТИТУЛЬНАЯ СТРАНИЦА (по шаблону title_lab.docx) =====
add_centered_text('Министерство науки и высшего образования Российской Федерации', size=14)
add_centered_text('Санкт-Петербургский политехнический университет Петра Великого', size=14)
add_empty_line()
add_centered_text('Институт компьютерных наук и кибербезопасности', size=14)
add_centered_text('Высшая школа кибербезопасности', size=14)

for _ in range(4):
    add_empty_line()

add_centered_text('ЛАБОРАТОРНАЯ РАБОТА №1', bold=True, size=16)
add_centered_text('«Управление ФС/реестром посредством CLI-приложения\nна ЯП Python»', size=14)
add_empty_line()
add_centered_text('по дисциплине', size=14)
add_centered_text('«Цифровая аналитика»', size=14)

for _ in range(4):
    add_empty_line()

# Выполнил (правая часть)
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
# Используем табуляцию для выравнивания
run = p.add_run('Выполнил')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('студент гр. __________\t\t\tВ.С. Тоцкий')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Преподаватель')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('__________________________\t\t\t____________')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

for _ in range(3):
    add_empty_line()

add_centered_text('Санкт-Петербург – 2026', size=14)

doc.add_page_break()

# ===== РАЗДЕЛ 1. ЦЕЛЬ РАБОТЫ =====
add_heading_text('1. Цель работы')
add_empty_line()

add_body_text(
    'Получение навыков управления файловой системой и реестром '
    'с использованием языка программирования Python.'
)

doc.add_page_break()

# ===== РАЗДЕЛ 2. ЗАДАНИЕ НА РАБОТУ =====
add_heading_text('2. Задание на работу')
add_empty_line()

tasks = [
    'Установить Python 3.X на систему.',
    'Установить IDE для разработки с помощью языка программирования Python.',
    'Реализовать консольное приложение на языке программирования Python, которое будет поддерживать '
    'операции управления ФС: создание файла (f_create), удаление файла (f_delete), запись в файл (f_write), '
    'чтение из файла (f_read), копирование файла (f_copy), переименование файла (f_rename).',
    'Для обработки входных параметров использовать модуль argparse. Для работы с ФС могут быть '
    'использованы модули: os, shutil.',
    'Обеспечить корректную обработку исключений с форматами сообщений «[+]» и «[-]».',
    'Написать отчёт о проделанной работе.',
    'Ответить на контрольные вопросы.',
]
for i, task in enumerate(tasks, 1):
    add_body_text(f'{i}. {task}')

doc.add_page_break()

# ===== РАЗДЕЛ 3. ХОД РАБОТЫ =====
add_heading_text('3. Ход работы')
add_empty_line()

add_body_text(
    'В ходе выполнения лабораторной работы было разработано консольное приложение '
    'на языке Python для управления файловой системой. Приложение реализовано в файле '
    'lab1.py и поддерживает следующие операции:'
)

functions_desc = [
    ('f_create(path)', 'создание файла по указанному пути. При необходимости создаются '
     'промежуточные директории с помощью os.makedirs.'),
    ('f_delete(path)', 'удаление файла по указанному пути. Проверяется существование файла '
     'и то, что путь не указывает на директорию.'),
    ('f_write(path, content)', 'запись содержимого в файл. Если файл или промежуточные '
     'директории не существуют, они создаются автоматически.'),
    ('f_read(path)', 'чтение содержимого файла и вывод его в консоль. Функция возвращает '
     'прочитанное содержимое.'),
    ('f_copy(src, dest)', 'копирование файла из одной директории в другую с помощью '
     'shutil.copy2. При необходимости создаются промежуточные директории назначения.'),
    ('f_rename(src, dest)', 'переименование (перемещение) файла с помощью os.rename. '
     'При необходимости создаются промежуточные директории назначения.'),
]

for fname, desc in functions_desc:
    p = doc.add_paragraph()
    run = p.add_run(f'{fname}')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run = p.add_run(f' — {desc}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

add_empty_line()

add_body_text(
    'Для обработки аргументов командной строки используется модуль argparse '
    'с подкомандами (subparsers) для каждой операции. Каждая подкоманда принимает '
    'соответствующие аргументы.'
)

add_empty_line()

add_body_text(
    'Все функции реализуют комплексную обработку исключений. Валидация входных '
    'параметров (None, пустая строка, пробелы, не-строковый тип) осуществляется '
    'вспомогательной функцией _validate_path с принудительным вызовом ValueError '
    'через инструкцию raise. Обрабатываются исключения FileNotFoundError, '
    'IsADirectoryError, PermissionError, UnicodeEncodeError, OSError и другие. '
    'Сообщения об ошибках имеют формат «[-]», сообщения об успешном выполнении — «[+]». '
    'Каждое сообщение об ошибке содержит значение аргумента, вызвавшего исключение. '
    'После вывода сообщения об ошибке исключение пробрасывается далее (re-raise).'
)

add_empty_line()

add_body_text('Используемые модули:')

modules = [
    ('argparse', 'обработка аргументов командной строки;'),
    ('os', 'операции с файловой системой (создание, удаление, переименование файлов и директорий);'),
    ('shutil', 'высокоуровневые операции с файлами (копирование).'),
]
for mod, desc in modules:
    p = doc.add_paragraph()
    run = p.add_run(mod)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run = p.add_run(f' — {desc}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

add_empty_line()

add_body_text('Примеры использования приложения:')
add_empty_line()

examples = [
    'python lab1.py create test.txt',
    'python lab1.py write test.txt "Hello, World!"',
    'python lab1.py read test.txt',
    'python lab1.py copy test.txt backup/test.txt',
    'python lab1.py rename test.txt new_test.txt',
    'python lab1.py delete new_test.txt',
]

for ex in examples:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(ex)
    run.font.name = 'Courier New'
    run.font.size = Pt(12)

add_empty_line()

# Листинг кода
add_body_text('В листинге 1 представлен исходный код приложения lab1.py.')
add_empty_line()

code_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'lab1.py')
with open(code_path, 'r', encoding='utf-8') as f:
    code = f.read()

# Заменить 4-пробельные отступы на 2-пробельные
lines = code.split('\n')
new_lines = []
for line in lines:
    spaces = len(line) - len(line.lstrip(' '))
    new_indent = ' ' * (spaces // 2)
    new_lines.append(new_indent + line.lstrip(' '))
code = '\n'.join(new_lines)

add_listing(code, 'Листинг 1 — Исходный код lab1.py')

doc.add_page_break()

# ===== РАЗДЕЛ 4. ОТВЕТЫ НА КОНТРОЛЬНЫЕ ВОПРОСЫ =====
add_heading_text('4. Ответы на контрольные вопросы')
add_empty_line()

# Вопрос 1
p = doc.add_paragraph()
run = p.add_run('1) Чем компилируемый язык программирования отличается от интерпретируемого? '
                'Приведите примеры языков программирования для обоих типов.')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_empty_line()

add_body_text(
    'Компилируемый язык программирования транслирует весь исходный код в машинный код '
    '(или промежуточный байт-код) до начала выполнения программы с помощью компилятора. '
    'Результатом является исполняемый файл, который может запускаться без исходного кода. '
    'Примеры: C, C++, Rust, Go.'
)

add_body_text(
    'Интерпретируемый язык выполняет программу построчно с помощью интерпретатора, '
    'который читает и исполняет инструкции непосредственно во время работы программы. '
    'Исходный код необходим для каждого запуска. '
    'Примеры: Python, JavaScript, Ruby, PHP.'
)

add_body_text(
    'Основные различия: компилируемые языки, как правило, обеспечивают более высокую '
    'производительность, но требуют этапа компиляции; интерпретируемые языки обеспечивают '
    'более быструю разработку и отладку, но могут работать медленнее.'
)

add_empty_line()

# Вопрос 2
p = doc.add_paragraph()
run = p.add_run('2) Каким образом могут быть считаны входные данные для консольного '
                'приложения помимо использования модуля argparse?')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_empty_line()

add_body_text(
    'Помимо модуля argparse, входные данные для консольного приложения могут быть считаны '
    'следующими способами:'
)

methods = [
    'функция input() — интерактивный ввод данных от пользователя;',
    'модуль sys (sys.argv) — прямой доступ к аргументам командной строки в виде списка строк;',
    'переменные окружения (os.environ) — чтение настроек из переменных окружения;',
    'чтение из файла конфигурации (json, yaml, ini) — загрузка параметров из конфигурационных файлов;',
    'стандартный поток ввода (sys.stdin) — чтение данных из перенаправленного ввода (pipe);',
    'модуль click — сторонняя библиотека для создания CLI-интерфейсов.',
]
for m in methods:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'— {m}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

add_empty_line()

# Вопрос 3
p = doc.add_paragraph()
run = p.add_run('3) В чём разница открытия файла конструкцией fd = open(...) / fd.close() '
                'от конструкции with open(...) as fd?')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_empty_line()

add_body_text(
    'При использовании конструкции fd = open("test.txt", "r") и последующего fd.close() '
    'программист самостоятельно управляет закрытием файла. Если между открытием и закрытием '
    'возникнет исключение, файл может остаться незакрытым, что приведёт к утечке ресурсов.'
)

add_body_text(
    'Конструкция with open("test.txt", "r") as fd является контекстным менеджером. '
    'Она гарантирует автоматическое закрытие файла при выходе из блока with, даже если '
    'внутри блока возникнет исключение. Это более безопасный и рекомендуемый способ '
    'работы с файлами в Python.'
)

add_empty_line()

# Вопрос 4
p = doc.add_paragraph()
run = p.add_run('4) Каким образом могут быть установлены модули для Python?')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_empty_line()

add_body_text('Модули для Python могут быть установлены следующими способами:')

install_methods = [
    'pip install <имя_модуля> — стандартный менеджер пакетов Python, устанавливает пакеты из PyPI;',
    'pip install <файл>.whl — установка из локального файла формата wheel;',
    'pip install <архив>.tar.gz — установка из архива с исходным кодом;',
    'python setup.py install — установка из исходного кода с помощью скрипта setup.py;',
    'conda install <имя_модуля> — менеджер пакетов Anaconda;',
    'ручное копирование модуля в директорию site-packages или в директорию проекта.',
]
for m in install_methods:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'— {m}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

doc.add_page_break()

# ===== РАЗДЕЛ 5. ВЫВОД =====
add_heading_text('5. Вывод')
add_empty_line()

add_body_text(
    'В ходе выполнения лабораторной работы было разработано консольное приложение '
    'на языке Python для управления файловой системой. Приложение поддерживает операции '
    'создания, удаления, записи, чтения, копирования и переименования файлов. '
    'Реализована комплексная обработка исключений с информативными сообщениями об ошибках '
    'в формате «[+]»/«[-]». Для обработки аргументов командной строки использован модуль '
    'argparse. Приложение успешно прошло автоматическое тестирование (199 из 199 тестов).'
)

# --- Сохранение ---
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Отчёт_ЛР1_Тоцкий_ВС.docx')
doc.save(output_path)
print(f'Report saved: {output_path}')
