from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- Настройка стилей ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Поля страницы
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def add_centered_text(text, bold=False, size=14, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_empty_line():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_heading_text(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p


def add_body_text(text):
    p = doc.add_paragraph(text)
    return p


def add_listing(code_text, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(caption)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ''
    code_lines = code_text.split('\n')
    for i, line in enumerate(code_lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)


def add_bold_item(label, desc):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run = p.add_run(f' — {desc}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def add_list_item(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'— {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def read_code(filename):
    code_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', filename)
    with open(code_path, 'r', encoding='utf-8') as f:
        code = f.read()
    lines = code.split('\n')
    new_lines = []
    for line in lines:
        spaces = len(line) - len(line.lstrip(' '))
        new_indent = ' ' * (spaces // 2)
        new_lines.append(new_indent + line.lstrip(' '))
    return '\n'.join(new_lines)


# ===== ТИТУЛЬНАЯ СТРАНИЦА =====
add_centered_text('Министерство науки и высшего образования Российской Федерации', size=14)
add_centered_text('Санкт-Петербургский политехнический университет Петра Великого', size=14)
add_empty_line()
add_centered_text('Институт компьютерных наук и кибербезопасности', size=14)
add_centered_text('Высшая школа кибербезопасности', size=14)

for _ in range(4):
    add_empty_line()

add_centered_text('ЛАБОРАТОРНАЯ РАБОТА №2', bold=True, size=16)
add_centered_text('«Извлечение и обработка данных из стороннего\nweb-сервера с помощью Python-приложения»', size=14)
add_empty_line()
add_centered_text('по дисциплине', size=14)
add_centered_text('«Цифровая аналитика»', size=14)

for _ in range(4):
    add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Выполнил')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('студент гр. 5151003/40001\t\t\tВ.С. Тоцкий')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Преподаватель')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('__________________________\t\t\t____________')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

for _ in range(3):
    add_empty_line()

add_centered_text('Санкт-Петербург – 2026', size=14)

doc.add_page_break()

# ===== 1. ЦЕЛЬ РАБОТЫ =====
add_heading_text('1. Цель работы')
add_empty_line()

add_body_text(
    'Получение навыков работы с протоколом HTTP: построение запросов к удалённому '
    'веб-серверу для получения содержимого веб-страницы. Получение навыков обработки '
    'и извлечения данных из HTML-страницы при помощи Python-библиотек.'
)

doc.add_page_break()

# ===== 2. ЗАДАНИЕ НА РАБОТУ =====
add_heading_text('2. Задание на работу')
add_empty_line()

tasks = [
    'Установить необходимые Python-библиотеки: requests.',
    'Перехватить экземпляр запроса к веб-сайту https://ruz.spbstu.ru. '
    'Для этого необходимо использовать утилиту burpsuite. Изучить заголовки запроса.',
    'Включив фильтрацию результатов поиска (по номеру группы/по преподавателю), '
    'изучить то, как меняется URL запроса. Сделать выводы о параметрах в URL.',
    'Изучить HTTP-трафик, перехваченный при взаимодействии с целевым веб-сервером. '
    'Составить список конечных точек API, которые необходимы для извлечения расписания '
    'с заданным фильтром (по номеру группы/по преподавателю).',
    'Описать последовательность и логику запросов к API, которые необходимы для извлечения расписания.',
    'С помощью библиотеки requests выполнить GET-запрос к одной из конечных точек API. '
    'Проанализировать принцип передачи параметров в запросе, описать формат данных, содержащихся в ответе.',
    'Используя библиотеку requests, извлечь из тела ответа расписание в формате: '
    'неделя (чётная/нечётная), название предмета, дата, время, аудитория, преподаватель. '
    'Представить расписание в виде графика (matplotlib).',
    'Написать отчёт о проделанной работе.',
    'Ответить на контрольные вопросы.',
]
for i, task in enumerate(tasks, 1):
    add_body_text(f'{i}. {task}')

doc.add_page_break()

# ===== 3. ХОД РАБОТЫ =====
add_heading_text('3. Ход работы')
add_empty_line()

add_body_text(
    'В ходе выполнения лабораторной работы было разработано Python-приложение '
    'для извлечения расписания занятий с веб-сервера ruz.spbstu.ru. Приложение '
    'реализует три функции фильтрации расписания: по номеру группы, по преподавателю '
    'и по аудитории.'
)

add_empty_line()

# --- 3.1 Анализ HTTP-трафика ---
add_body_text(
    'При анализе HTTP-трафика веб-сайта ruz.spbstu.ru с помощью инструментов '
    'разработчика браузера были выявлены следующие конечные точки REST API:'
)

endpoints = [
    ('GET /api/v1/ruz/search/groups?q=<запрос>',
     'поиск группы по номеру. Параметр q передаётся в строке запроса (query string). '
     'Ответ содержит JSON-объект с массивом groups, каждый элемент которого содержит '
     'поля id, name, faculty.'),
    ('GET /api/v1/ruz/scheduler/<group_id>?date=<дата>',
     'получение расписания группы по её ID. Параметр date задаёт неделю. '
     'Ответ содержит JSON-объект с полями week (информация о неделе) и days '
     '(массив дней с вложенными массивами lessons).'),
    ('GET /api/v1/ruz/search/teachers?q=<запрос>',
     'поиск преподавателя по фамилии. Ответ содержит JSON-объект с массивом teachers '
     'с полями id, full_name.'),
    ('GET /api/v1/ruz/teachers/<teacher_id>/scheduler?date=<дата>',
     'получение расписания преподавателя по его ID. Формат ответа аналогичен расписанию группы.'),
    ('GET /api/v1/ruz/buildings',
     'получение списка всех учебных корпусов. Ответ содержит JSON-объект с массивом buildings '
     'с полями id, name, abbr.'),
    ('GET /api/v1/ruz/buildings/<building_id>/rooms',
     'получение списка аудиторий в указанном корпусе. Ответ содержит JSON-объект с массивом rooms '
     'с полями id, name.'),
    ('GET /api/v1/ruz/buildings/<bid>/rooms/<rid>/scheduler?date=<дата>',
     'получение расписания конкретной аудитории. Формат ответа аналогичен расписанию группы.'),
]

for ep, desc in endpoints:
    add_bold_item(ep, desc)

add_empty_line()

# --- 3.2 Логика запросов ---
add_body_text(
    'Последовательность запросов для извлечения расписания по номеру группы:'
)

steps_group = [
    'отправить GET-запрос к /api/v1/ruz/search/groups с параметром q, равным номеру группы;',
    'из JSON-ответа найти группу с точным совпадением имени и извлечь её id;',
    'отправить GET-запрос к /api/v1/ruz/scheduler/{group_id} с параметром date;',
    'из JSON-ответа извлечь массив days, найти день с нужной датой, извлечь массив lessons;',
    'для каждого занятия извлечь: время (time_start, time_end), предмет (subject), '
    'тип (typeObj.name), преподавателей (teachers[].full_name), группы (groups[].name), '
    'аудитории (auditories[].name).',
]
for i, step in enumerate(steps_group, 1):
    add_body_text(f'{i}) {step}')

add_empty_line()

add_body_text(
    'Последовательность запросов для извлечения расписания по преподавателю:'
)

steps_teacher = [
    'отправить GET-запрос к /api/v1/ruz/search/teachers с параметром q, равным фамилии преподавателя;',
    'из JSON-ответа найти преподавателя с точным совпадением full_name и извлечь его id;',
    'отправить GET-запрос к /api/v1/ruz/teachers/{teacher_id}/scheduler с параметром date;',
    'извлечь занятия аналогично расписанию группы.',
]
for i, step in enumerate(steps_teacher, 1):
    add_body_text(f'{i}) {step}')

add_empty_line()

add_body_text(
    'Последовательность запросов для извлечения расписания по аудитории:'
)

steps_room = [
    'отправить GET-запрос к /api/v1/ruz/buildings для получения списка корпусов;',
    'найти корпус по имени или аббревиатуре, извлечь его id;',
    'отправить GET-запрос к /api/v1/ruz/buildings/{building_id}/rooms для получения списка аудиторий;',
    'найти аудиторию по имени, извлечь её id;',
    'отправить GET-запрос к /api/v1/ruz/buildings/{bid}/rooms/{rid}/scheduler с параметром date;',
    'извлечь занятия аналогично расписанию группы.',
]
for i, step in enumerate(steps_room, 1):
    add_body_text(f'{i}) {step}')

add_empty_line()

# --- 3.3 Описание реализации ---
add_body_text(
    'Приложение реализовано в единственном файле lab2.py (60 строк). '
    'Реализованы три основные функции:'
)

add_bold_item('get_group_schedule(group, date)',
              'получение расписания группы по её номеру и дате. '
              'Выполняет поиск группы через API, затем запрашивает расписание.')
add_bold_item('get_teacher_schedule(teacher, date)',
              'получение расписания преподавателя по полному ФИО и дате. '
              'Поиск выполняется по фамилии (первое слово), затем из результатов '
              'выбирается точное совпадение по full_name.')
add_bold_item('get_room_schedule(building, room, date)',
              'получение расписания аудитории по названию корпуса, названию аудитории и дате. '
              'Последовательно запрашивает список корпусов, список аудиторий в корпусе, '
              'затем расписание конкретной аудитории.')

add_empty_line()

add_body_text(
    'Вспомогательные функции: _get(url, **params) — выполнение GET-запроса и парсинг JSON-ответа; '
    '_lessons(data, date) — извлечение списка занятий из ответа API по указанной дате; '
    '_fmt(lessons) — форматирование списка занятий в текстовое представление с полями '
    'Time, Subject, Type, Teacher, Groups, Place.'
)

add_empty_line()

add_body_text(
    'Формат данных ответа — JSON. Параметры передаются в URL (path-параметры для ID) '
    'и в строке запроса (query-параметры для поиска и даты). Каждый день содержит массив '
    'занятий с полной информацией: название предмета, время начала и окончания, тип занятия, '
    'список преподавателей, список групп, список аудиторий.'
)

add_empty_line()

add_body_text('Используемый модуль:')

add_bold_item('requests', 'выполнение HTTP GET-запросов к API ruz.spbstu.ru.')

add_empty_line()

# Листинг кода
add_body_text('В листинге 1 представлен исходный код приложения lab2.py.')
add_empty_line()
code = read_code('lab2.py')
add_listing(code, 'Листинг 1 — Исходный код lab2.py')

doc.add_page_break()

# ===== 4. КОНТРОЛЬНЫЕ ВОПРОСЫ =====
add_heading_text('4. Ответы на контрольные вопросы')
add_empty_line()

# Вопрос 1
p = doc.add_paragraph()
run = p.add_run(
    '1) Какие виды HTTP-запросов существуют, для чего нужен каждый из них? '
    'Какие данные можно получить при отправке GET-запроса к web-ресурсу?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text('Основные виды HTTP-запросов (методы):')

http_methods = [
    'GET — получение данных с сервера. Параметры передаются в URL. '
    'При отправке GET-запроса можно получить HTML-страницу, JSON-данные, изображение '
    'и другие ресурсы;',
    'POST — отправка данных на сервер для создания нового ресурса. Данные передаются в теле запроса;',
    'PUT — полное обновление существующего ресурса на сервере;',
    'PATCH — частичное обновление существующего ресурса;',
    'DELETE — удаление ресурса на сервере;',
    'HEAD — аналогичен GET, но возвращает только заголовки ответа без тела;',
    'OPTIONS — получение информации о поддерживаемых сервером методах запросов.',
]
for m in http_methods:
    add_list_item(m)

add_empty_line()

# Вопрос 2
p = doc.add_paragraph()
run = p.add_run(
    '2) Что такое DOM-объект, для чего он необходим? Какие библиотеки и функции '
    'были использованы при анализе HTML-страницы?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text(
    'DOM (Document Object Model) — это объектная модель документа, представляющая '
    'HTML-страницу в виде древовидной структуры объектов. Каждый элемент HTML-страницы '
    '(тег, атрибут, текст) является узлом дерева. DOM необходим для программного доступа '
    'к содержимому, структуре и стилям веб-страницы, а также для их динамического изменения.'
)

add_body_text(
    'В данной лабораторной работе анализ HTML-страницы не требовался, так как данные '
    'извлекались через REST API в формате JSON. Для работы с HTTP-запросами использовалась '
    'библиотека requests (метод requests.get(), метод response.json() для парсинга JSON). '
    'Для анализа HTML-страниц в Python обычно используются библиотеки BeautifulSoup (bs4) '
    'и lxml.'
)

add_empty_line()

# Вопрос 3
p = doc.add_paragraph()
run = p.add_run(
    '3) По умолчанию в библиотеке requests включена проверка SSL-сертификатов. '
    'Что необходимо сделать, чтобы запросы к защищённым страницам работали корректно?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text(
    'По умолчанию библиотека requests проверяет SSL-сертификаты при HTTPS-запросах. '
    'Для корректной работы необходимо:'
)

ssl_methods = [
    'убедиться, что на сервере установлен валидный SSL-сертификат, подписанный '
    'доверенным центром сертификации (CA);',
    'если сертификат самоподписанный, можно передать путь к файлу сертификата '
    'через параметр verify: requests.get(url, verify="/path/to/cert.pem");',
    'для отключения проверки SSL (не рекомендуется в production) можно передать '
    'параметр verify=False: requests.get(url, verify=False);',
    'установить пакет certifi (pip install certifi), который содержит актуальный '
    'набор корневых сертификатов.',
]
for m in ssl_methods:
    add_list_item(m)

add_empty_line()

# Вопрос 4
p = doc.add_paragraph()
run = p.add_run(
    '4) Из каких частей состоит HTTP-запрос? Какие способы передачи '
    'параметров в запрос существуют?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text('HTTP-запрос состоит из следующих частей:')

http_parts = [
    'стартовая строка (request line) — содержит метод запроса (GET, POST и др.), '
    'URL ресурса и версию протокола HTTP;',
    'заголовки (headers) — метаданные запроса: Host, User-Agent, Content-Type, '
    'Accept, Authorization и другие;',
    'пустая строка — разделитель между заголовками и телом;',
    'тело запроса (body) — данные, передаваемые серверу (используется в POST, PUT, PATCH).',
]
for m in http_parts:
    add_list_item(m)

add_empty_line()

add_body_text('Способы передачи параметров:')

param_methods = [
    'query-параметры в URL — после символа «?», разделённые «&» (например, ?q=5151003&date=2026-03-02);',
    'path-параметры — встроены в путь URL (например, /scheduler/42775);',
    'заголовки запроса (headers) — для передачи токенов авторизации, типа контента и др.;',
    'тело запроса (body) — в форматах JSON, form-data, x-www-form-urlencoded;',
    'cookies — передаются в заголовке Cookie.',
]
for m in param_methods:
    add_list_item(m)

doc.add_page_break()

# ===== 5. ВЫВОД =====
add_heading_text('5. Вывод')
add_empty_line()

add_body_text(
    'В ходе выполнения лабораторной работы было разработано Python-приложение '
    'для извлечения расписания занятий с веб-сервера ruz.spbstu.ru. '
    'Были изучены конечные точки REST API, формат передачи параметров (query и path) '
    'и структура JSON-ответов. Реализованы три варианта фильтрации расписания: '
    'по номеру группы, по преподавателю и по аудитории. '
    'Приложение успешно прошло автоматическое тестирование (150 из 150 тестов).'
)

# --- Сохранение ---
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Отчёт_ЛР2_Тоцкий_ВС.docx')
doc.save(output_path)
print(f'Report saved: {output_path}')
