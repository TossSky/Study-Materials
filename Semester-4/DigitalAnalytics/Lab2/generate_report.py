from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- Настройка стилей ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Поля страницы
for section in doc.sections:
  section.top_margin = Cm(2)
  section.bottom_margin = Cm(2)
  section.left_margin = Cm(2)
  section.right_margin = Cm(2)


def add_centered_text(text, bold=False, size=14, space_after=0):
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p.paragraph_format.first_line_indent = Cm(0)
  p.paragraph_format.space_after = Pt(space_after)
  run = p.add_run(text)
  run.bold = bold
  run.font.size = Pt(size)
  run.font.name = 'Times New Roman'
  return p


def add_empty_line():
  p = doc.add_paragraph()
  p.paragraph_format.first_line_indent = Cm(0)
  return p


def add_heading_text(text):
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p.paragraph_format.first_line_indent = Cm(0)
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(0)
  run = p.add_run(text)
  run.bold = True
  run.font.size = Pt(14)
  run.font.name = 'Times New Roman'
  return p


def add_body_text(text):
  p = doc.add_paragraph(text)
  return p


def add_listing(code_text, caption):
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p.paragraph_format.first_line_indent = Cm(0)
  run = p.add_run(caption)
  run.font.size = Pt(14)
  run.font.name = 'Times New Roman'

  table = doc.add_table(rows=1, cols=1)
  table.style = 'Table Grid'
  cell = table.cell(0, 0)
  cell.text = ''
  code_lines = code_text.split('\n')
  for i, line in enumerate(code_lines):
    if i == 0:
      p = cell.paragraphs[0]
    else:
      p = cell.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(line)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)


def read_code(filename):
  code_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
  with open(code_path, 'r', encoding='utf-8') as f:
    return f.read()


# ===== ТИТУЛЬНАЯ СТРАНИЦА =====
add_centered_text('Министерство науки и высшего образования Российской Федерации', size=14)
add_centered_text('Санкт-Петербургский политехнический университет Петра Великого', size=14)
add_empty_line()
add_centered_text('Институт компьютерных наук и кибербезопасности', size=14)
add_centered_text('Высшая школа кибербезопасности', size=14)

for _ in range(4):
  add_empty_line()

add_centered_text('ЛАБОРАТОРНАЯ РАБОТА №2', bold=True, size=16)
add_centered_text('«Извлечение и обработка данных из стороннего\nweb-сервера с помощью Python-приложения»', size=14)
add_empty_line()
add_centered_text('по дисциплине', size=14)
add_centered_text('«Цифровая аналитика»', size=14)

for _ in range(4):
  add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Выполнил')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('студент гр. __________\t\t\tВ.С. Тоцкий')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Преподаватель')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('__________________________\t\t\t____________')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

for _ in range(3):
  add_empty_line()

add_centered_text('Санкт-Петербург – 2026', size=14)

doc.add_page_break()

# ===== 1. ЦЕЛЬ РАБОТЫ =====
add_heading_text('1. Цель работы')
add_empty_line()

add_body_text(
  'Получение навыков работы с протоколом HTTP: построение запросов к удалённому '
  'веб-серверу для получения содержимого веб-страницы. Получение навыков обработки '
  'и извлечения данных из HTML-страницы при помощи Python-библиотек.'
)

doc.add_page_break()

# ===== 2. ЗАДАНИЕ НА РАБОТУ =====
add_heading_text('2. Задание на работу')
add_empty_line()

tasks = [
  'Установить необходимые Python-библиотеки: requests.',
  'Перехватить экземпляр запроса к веб-сайту https://ruz.spbstu.ru. Изучить заголовки запроса.',
  'Включив фильтрацию результатов поиска, изучить то, как меняется URL запроса.',
  'Изучить HTTP-трафик, перехваченный при взаимодействии с целевым веб-сервером. '
  'Составить список конечных точек API.',
  'Описать последовательность и логику запросов к API для извлечения расписания.',
  'С помощью библиотеки requests выполнить GET-запрос к одной из конечных точек API. '
  'Проанализировать принцип передачи параметров, описать формат данных в ответе.',
  'Используя библиотеку requests, извлечь из тела ответа расписание в формате: '
  'неделя (чётная/нечётная), название предмета, дата, время, аудитория, преподаватель. '
  'Представить расписание в виде графика (matplotlib).',
  'Написать отчёт о проделанной работе.',
  'Ответить на контрольные вопросы.',
]
for i, task in enumerate(tasks, 1):
  add_body_text(f'{i}. {task}')

doc.add_page_break()

# ===== 3. ХОД РАБОТЫ =====
add_heading_text('3. Ход работы')
add_empty_line()

add_body_text(
  'В ходе выполнения лабораторной работы было разработано консольное приложение '
  'на языке Python для извлечения расписания занятий с веб-сервера ruz.spbstu.ru.'
)

add_empty_line()

add_body_text(
  'При анализе HTTP-трафика веб-сайта ruz.spbstu.ru были выявлены следующие '
  'конечные точки API, необходимые для извлечения расписания:'
)

endpoints = [
  ('GET /api/v1/ruz/search/groups?q=<запрос>', 'поиск группы по номеру. '
   'Параметр q передаётся в строке запроса (query string). Ответ содержит JSON-массив '
   'найденных групп с полями id, name, faculty.'),
  ('GET /api/v1/ruz/scheduler/<group_id>', 'получение расписания группы по её ID. '
   'Ответ содержит JSON-объект с полями week (информация о неделе), days (массив дней '
   'с вложенными массивами lessons), group (информация о группе).'),
  ('GET /api/v1/ruz/search/teachers?q=<запрос>', 'поиск преподавателя по ФИО. '
   'Параметр q передаётся в строке запроса. Ответ содержит JSON-массив найденных '
   'преподавателей с полями id, full_name, chair.'),
  ('GET /api/v1/ruz/teachers/<teacher_id>/scheduler', 'получение расписания преподавателя '
   'по его ID. Формат ответа аналогичен расписанию группы.'),
]

for ep, desc in endpoints:
  p = doc.add_paragraph()
  run = p.add_run(ep)
  run.bold = True
  run.font.name = 'Times New Roman'
  run.font.size = Pt(14)
  run = p.add_run(f' — {desc}')
  run.font.name = 'Times New Roman'
  run.font.size = Pt(14)

add_empty_line()

add_body_text(
  'Последовательность запросов для извлечения расписания по номеру группы:'
)

steps = [
  'отправить GET-запрос к /api/v1/ruz/search/groups с параметром q, равным номеру группы;',
  'из JSON-ответа извлечь поле id первой найденной группы;',
  'отправить GET-запрос к /api/v1/ruz/scheduler/{group_id};',
  'из JSON-ответа извлечь данные о неделе (week) и массив дней (days) с занятиями (lessons);',
  'для каждого занятия извлечь: название предмета (subject), время (time_start, time_end), '
  'тип занятия (typeObj.name), преподавателя (teachers[].full_name), '
  'аудиторию (auditories[].name, auditories[].building.abbr).',
]
for i, step in enumerate(steps, 1):
  add_body_text(f'{i}) {step}')

add_empty_line()

add_body_text(
  'Формат данных ответа — JSON. Параметры передаются в URL (path-параметры для ID) '
  'и в строке запроса (query-параметры для поиска). Каждый день содержит массив занятий '
  'с полной информацией: название предмета, время начала и окончания, тип занятия, '
  'список преподавателей, список аудиторий с указанием корпуса, чётность недели (parity: '
  '0 — обе недели, 1 — нечётная, 2 — чётная).'
)

add_empty_line()

add_body_text(
  'Помимо текстового вывода расписания, приложение строит столбчатую диаграмму '
  'с использованием библиотеки matplotlib. По оси абсцисс — день недели, по оси '
  'ординат — количество занятий в этот день.'
)

add_empty_line()

add_body_text('Используемые модули:')

modules = [
  ('requests', 'выполнение HTTP-запросов к API ruz.spbstu.ru;'),
  ('matplotlib.pyplot', 'построение столбчатой диаграммы расписания;'),
  ('argparse', 'обработка аргументов командной строки.'),
]
for mod, desc in modules:
  p = doc.add_paragraph()
  run = p.add_run(mod)
  run.bold = True
  run.font.name = 'Times New Roman'
  run.font.size = Pt(14)
  run = p.add_run(f' — {desc}')
  run.font.name = 'Times New Roman'
  run.font.size = Pt(14)

add_empty_line()

add_body_text('Примеры использования приложения:')
add_empty_line()

examples = [
  'python lab2.py "5151003/30802"',
  'python lab2_v2.py "Семьянов"',
]
for ex in examples:
  p = doc.add_paragraph()
  p.paragraph_format.first_line_indent = Cm(0)
  p.paragraph_format.left_indent = Cm(1.25)
  run = p.add_run(ex)
  run.font.name = 'Courier New'
  run.font.size = Pt(12)

add_empty_line()

# Листинги кода
add_body_text('В листинге 1 представлен исходный код приложения lab2.py (вариант 1 — фильтрация по группе).')
add_empty_line()
code1 = read_code('lab2.py')
add_listing(code1, 'Листинг 1 — Исходный код lab2.py')

add_empty_line()

add_body_text('В листинге 2 представлен исходный код приложения lab2_v2.py (вариант 2 — фильтрация по преподавателю).')
add_empty_line()
code2 = read_code('lab2_v2.py')
add_listing(code2, 'Листинг 2 — Исходный код lab2_v2.py')

doc.add_page_break()

# ===== 4. КОНТРОЛЬНЫЕ ВОПРОСЫ =====
add_heading_text('4. Ответы на контрольные вопросы')
add_empty_line()

# Вопрос 1
p = doc.add_paragraph()
run = p.add_run(
  '1) Какие виды HTTP-запросов существуют, для чего нужен каждый из них? '
  'Какие данные можно получить при отправке GET-запроса к web-ресурсу?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text(
  'Основные виды HTTP-запросов (методы):'
)

http_methods = [
  'GET — получение данных с сервера. Параметры передаются в URL. '
  'При отправке GET-запроса можно получить HTML-страницу, JSON-данные, изображение '
  'и другие ресурсы;',
  'POST — отправка данных на сервер для создания нового ресурса. Данные передаются в теле запроса;',
  'PUT — полное обновление существующего ресурса на сервере;',
  'PATCH — частичное обновление существующего ресурса;',
  'DELETE — удаление ресурса на сервере;',
  'HEAD — аналогичен GET, но возвращает только заголовки ответа без тела;',
  'OPTIONS — получение информации о поддерживаемых сервером методах запросов.',
]
for m in http_methods:
  p = doc.add_paragraph()
  p.paragraph_format.first_line_indent = Cm(0)
  p.paragraph_format.left_indent = Cm(1.25)
  run = p.add_run(f'— {m}')
  run.font.name = 'Times New Roman'
  run.font.size = Pt(14)

add_empty_line()

# Вопрос 2
p = doc.add_paragraph()
run = p.add_run(
  '2) Что такое DOM-объект, для чего он необходим? Какие библиотеки и функции '
  'были использованы при анализе HTML-страницы?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text(
  'DOM (Document Object Model) — это объектная модель документа, представляющая '
  'HTML-страницу в виде древовидной структуры объектов. Каждый элемент HTML-страницы '
  '(тег, атрибут, текст) является узлом дерева. DOM необходим для программного доступа '
  'к содержимому, структуре и стилям веб-страницы, а также для их динамического изменения.'
)

add_body_text(
  'В данной лабораторной работе анализ HTML-страницы не требовался, так как данные '
  'извлекались через REST API в формате JSON. Для работы с HTTP-запросами использовалась '
  'библиотека requests (метод requests.get(), метод response.json() для парсинга JSON). '
  'Для анализа HTML-страниц в Python обычно используются библиотеки BeautifulSoup (bs4) '
  'и lxml.'
)

add_empty_line()

# Вопрос 3
p = doc.add_paragraph()
run = p.add_run(
  '3) По умолчанию в библиотеке requests включена проверка SSL-сертификатов. '
  'Что необходимо сделать, чтобы запросы к защищённым страницам работали корректно?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text(
  'По умолчанию библиотека requests проверяет SSL-сертификаты при HTTPS-запросах. '
  'Для корректной работы необходимо:'
)

ssl_methods = [
  'убедиться, что на сервере установлен валидный SSL-сертификат, подписанный '
  'доверенным центром сертификации (CA);',
  'если сертификат самоподписанный, можно передать путь к файлу сертификата '
  'через параметр verify: requests.get(url, verify="/path/to/cert.pem");',
  'для отключения проверки SSL (не рекомендуется в production) можно передать '
  'параметр verify=False: requests.get(url, verify=False);',
  'установить пакет certifi (pip install certifi), который содержит актуальный '
  'набор корневых сертификатов.',
]
for m in ssl_methods:
  p = doc.add_paragraph()
  p.paragraph_format.first_line_indent = Cm(0)
  p.paragraph_format.left_indent = Cm(1.25)
  run = p.add_run(f'— {m}')
  run.font.name = 'Times New Roman'
  run.font.size = Pt(14)

add_empty_line()

# Вопрос 4
p = doc.add_paragraph()
run = p.add_run(
  '4) Из каких частей состоит HTTP-запрос? Какие способы передачи '
  'параметров в запрос существуют?'
)
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_empty_line()

add_body_text('HTTP-запрос состоит из следующих частей:')

http_parts = [
  'стартовая строка (request line) — содержит метод запроса (GET, POST и др.), '
  'URL ресурса и версию протокола HTTP;',
  'заголовки (headers) — метаданные запроса: Host, User-Agent, Content-Type, '
  'Accept, Authorization и другие;',
  'пустая строка — разделитель между заголовками и телом;',
  'тело запроса (body) — данные, передаваемые серверу (используется в POST, PUT, PATCH).',
]
for m in http_parts:
  p = doc.add_paragraph()
  p.paragraph_format.first_line_indent = Cm(0)
  p.paragraph_format.left_indent = Cm(1.25)
  run = p.add_run(f'— {m}')
  run.font.name = 'Times New Roman'
  run.font.size = Pt(14)

add_empty_line()

add_body_text('Способы передачи параметров:')

param_methods = [
  'query-параметры в URL — после символа «?», разделённые «&» (например, ?q=5151003&limit=10);',
  'path-параметры — встроены в путь URL (например, /scheduler/42811);',
  'заголовки запроса (headers) — для передачи токенов авторизации, типа контента и др.;',
  'тело запроса (body) — в форматах JSON, form-data, x-www-form-urlencoded;',
  'cookies — передаются в заголовке Cookie.',
]
for m in param_methods:
  p = doc.add_paragraph()
  p.paragraph_format.first_line_indent = Cm(0)
  p.paragraph_format.left_indent = Cm(1.25)
  run = p.add_run(f'— {m}')
  run.font.name = 'Times New Roman'
  run.font.size = Pt(14)

doc.add_page_break()

# ===== 5. ВЫВОД =====
add_heading_text('5. Вывод')
add_empty_line()

add_body_text(
  'В ходе выполнения лабораторной работы было разработано консольное приложение '
  'на языке Python для извлечения расписания занятий с веб-сервера ruz.spbstu.ru. '
  'Были изучены конечные точки API, формат передачи параметров (query и path) '
  'и структура JSON-ответов. Реализованы два варианта фильтрации: по номеру группы '
  'и по преподавателю. Расписание выводится в консоль в структурированном виде '
  'с указанием недели, предмета, даты, времени, аудитории и преподавателя. '
  'Дополнительно построена столбчатая диаграмма количества занятий по дням недели '
  'с использованием библиотеки matplotlib.'
)

# --- Сохранение ---
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Отчёт_ЛР2_Тоцкий_ВС.docx')
doc.save(output_path)
print(f'Report saved: {output_path}')
