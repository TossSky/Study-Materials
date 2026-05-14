from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- Настройка стилей ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def add_centered_text(text, bold=False, size=14, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_empty_line():
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_heading_text(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p


def add_body_text(text):
    p = doc.add_paragraph(text)
    return p


def add_listing(code_text, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(caption)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ''
    code_lines = code_text.split('\n')
    for i, line in enumerate(code_lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)


def add_bold_item(label, desc):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run = p.add_run(f' — {desc}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def add_list_item(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'— {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def read_code(filename):
    code_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', filename)
    with open(code_path, 'r', encoding='utf-8') as f:
        code = f.read()
    lines = code.split('\n')
    new_lines = []
    for line in lines:
        spaces = len(line) - len(line.lstrip(' '))
        new_indent = ' ' * (spaces // 2)
        new_lines.append(new_indent + line.lstrip(' '))
    return '\n'.join(new_lines)


# ===== ТИТУЛЬНАЯ СТРАНИЦА =====
add_centered_text('Министерство науки и высшего образования Российской Федерации', size=14)
add_centered_text('Санкт-Петербургский политехнический университет Петра Великого', size=14)
add_empty_line()
add_centered_text('Институт компьютерных наук и кибербезопасности', size=14)
add_centered_text('Высшая школа кибербезопасности', size=14)

for _ in range(4):
    add_empty_line()

add_centered_text('ЛАБОРАТОРНАЯ РАБОТА №4', bold=True, size=16)
add_centered_text('«Реализация Telegram-бота на языке программирования Python»', size=14)
add_empty_line()
add_centered_text('по дисциплине', size=14)
add_centered_text('«Цифровая аналитика»', size=14)

for _ in range(4):
    add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Выполнил')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('студент гр. 5151003/40001\t\t\t\t\t\t\tТоцкий В.')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

add_empty_line()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Преподаватель\t\t\t\t\t\t\t\t\tПисков А.А.')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

for _ in range(5):
    add_empty_line()

add_centered_text('Санкт-Петербург – 2026', size=14)

doc.add_page_break()

# ===== 1. ЦЕЛЬ РАБОТЫ =====
add_heading_text('1. Цель работы')
add_empty_line()

add_body_text(
    'Получение навыков создания Telegram-ботов на языке программирования Python. '
    'Изучение функционала webhook из Telegram Bot API, перенос ранее реализованного '
    'бота с polling на webhook, усовершенствование метода хранения пользовательских '
    'данных.'
)

doc.add_page_break()

# ===== 2. ЗАДАНИЕ НА РАБОТУ =====
add_heading_text('2. Задание на работу')
add_empty_line()

tasks = [
    'Выделить в текущей реализации Telegram-бота из лаб. 3 недостатки хранения '
    'пользовательских данных.',
    'Усовершенствовать метод хранения данных в соответствии с выделенными недостатками.',
    'Изучить функционал webhook из Telegram API (setWebhook).',
    'При помощи сервиса pythonanywhere.com создать веб-приложение, которое будет '
    'содержать конечные точки, выступающие в роли webhook’ов. В качестве фреймворка '
    'использовать Flask.',
    'Веб-приложение должно содержать обработчик POST-запросов. На данный обработчик '
    'в дальнейшем будут приходить запросы от Telegram API в формате JSON.',
    'Зарегистрировать созданный webhook в Telegram-боте при помощи Telegram API '
    '(setWebhook).',
    'Отправить сообщение в чат бота, использующего webhook. Ознакомиться с форматом '
    'полученного веб-приложением сообщения.',
    'Реализовать обработку получаемых сообщений и отправку ответных сообщений при '
    'помощи API sendMessage.',
    'Перенести функционал ранее созданного Telegram-бота из лаб. 3 на webhook’и.',
    'Выделить преимущества и недостатки реализаций Telegram-ботов с использованием '
    'polling’a и webhook’ов.',
]
for i, task in enumerate(tasks, 1):
    add_body_text(f'{i}. {task}')

doc.add_page_break()

# ===== 3. ХОД РАБОТЫ =====
add_heading_text('3. Ход работы')
add_empty_line()

add_body_text(
    'В ходе выполнения лабораторной работы был выполнен перенос ранее реализованного '
    'в лаб. 3 Telegram-бота с механизма long polling на webhook, а также '
    'усовершенствована схема хранения пользовательских данных. Веб-приложение '
    'реализовано на микрофреймворке Flask и развёрнуто на сервисе pythonanywhere.com. '
    'Функционал бота — регистрация, аутентификация и бинарная классификация изображений '
    'с использованием обученной в лаб. 3 модели ResNet-18 — полностью сохранён.'
)
add_empty_line()

# --- 3.1 Недостатки исходного хранения ---
add_body_text('3.1. Недостатки хранения пользовательских данных в реализации лаб. 3')
add_empty_line()

add_body_text(
    'Исходная реализация бота из лаб. 3 (файл bot_polling.py в репозитории) '
    'использует для хранения данных пользователей обычные словари Python в '
    'оперативной памяти процесса:'
)

drawbacks = [
    'отсутствие персистентности — все данные пользователей теряются при перезапуске '
    'процесса. Зарегистрированному пользователю приходится регистрироваться заново '
    'после каждой перезагрузки сервера;',
    'хранение паролей в открытом виде — поле password в словаре users[chat_id] '
    'содержит plaintext-строку. Любой, кто получит доступ к памяти процесса или '
    'к дампу (например, через ошибку, выгружающую stack trace), увидит пароль;',
    'отсутствие защиты от перебора и подбора — пароль сравнивается через оператор ==, '
    'который имеет ранний возврат и подвержен timing-атакам;',
    'race condition при многопоточной обработке — long polling в pyTelegramBotAPI '
    'по умолчанию однопоточный, но при переходе на webhook (WSGI) обработка может '
    'идти параллельно. Изменение users из нескольких потоков без блокировок '
    'приводит к потере обновлений;',
    'смешение состояния и данных — pending_action хранит «ожидаемое следующее '
    'сообщение» в той же структуре, что и пользовательские данные, без явной модели '
    'и без TTL. При сбое в логике пользователь может застрять в состоянии '
    '«ввод пароля» навсегда;',
    'отсутствие изоляции между chat_id — потенциальный баг кода или подмена '
    'идентификатора приводит к раскрытию данных другого пользователя.',
]
for d in drawbacks:
    add_list_item(d)

add_empty_line()

# --- 3.2 Усовершенствованное хранение ---
add_body_text('3.2. Усовершенствованная схема хранения')
add_empty_line()

add_body_text(
    'Хранение пользовательских данных вынесено в отдельный модуль storage.py, '
    'который инкапсулирует работу с базой данных и предоставляет публичный API: '
    'register, verify, set_logged_in, is_logged_in, set_pending, get_pending, '
    'clear_pending, exists. Принятые проектные решения:'
)

improvements = [
    'персистентное хранилище SQLite — стандартный модуль sqlite3 включён в Python, '
    'не требует отдельного сервера и поддерживает транзакции. База данных users.db '
    'создаётся при первом запуске функцией init_db;',
    'хеширование паролей PBKDF2-HMAC-SHA256 с уникальной 16-байтной солью на каждого '
    'пользователя и числом итераций 200 000 — соответствует современным рекомендациям '
    'NIST SP 800-132. Соль и хеш хранятся в полях password_salt и password_hash '
    'типа BLOB;',
    'сравнение хешей через secrets.compare_digest — выполняется за константное время '
    'и устойчиво к timing-атакам;',
    'потокобезопасность — все операции выполняются под глобальным threading.Lock в '
    'контекстном менеджере, а соединения SQLite открываются с PRAGMA journal_mode=WAL '
    'для повышения параллелизма чтения;',
    'выделение pending_action в отдельную колонку — состояние «ожидаемое следующее '
    'сообщение от пользователя» хранится в БД, а не в памяти, и переживает '
    'перезапуск процесса;',
    'идемпотентность регистрации — функция register возвращает False при попытке '
    'повторно зарегистрировать существующего пользователя, что предотвращает '
    'случайную перезапись пароля.',
]
for i in improvements:
    add_list_item(i)

add_empty_line()

# --- 3.3 Webhook ---
add_body_text('3.3. Изучение функционала webhook из Telegram Bot API')
add_empty_line()

add_body_text(
    'Telegram Bot API поддерживает два способа доставки обновлений приложению '
    'разработчика: long polling и webhook. Webhook — это HTTPS-эндпоинт, который '
    'разработчик регистрирует в Telegram с помощью метода setWebhook. После этого '
    'Telegram сам отправляет POST-запрос на указанный URL при каждом новом событии '
    '(сообщение, нажатие inline-кнопки, callback). Тело запроса — JSON-объект '
    'структуры Update.'
)

add_body_text(
    'Требования Telegram к webhook-эндпоинту:'
)

webhook_req = [
    'обязателен HTTPS с валидным TLS-сертификатом, выданным доверенным CA, либо '
    'самоподписанным сертификатом, передаваемым в параметре certificate;',
    'разрешены порты 443, 80, 88, 8443;',
    'эндпоинт должен возвращать 200 OK не позднее чем через 60 секунд после '
    'получения обновления, иначе Telegram повторит попытку;',
    'один URL — один бот; одновременно нельзя использовать polling и webhook на '
    'одном и том же боте;',
    'для защиты от подделок рекомендуется включать в путь URL секретный токен '
    '(например, /<BOT_TOKEN>) или использовать параметр secret_token.',
]
for w in webhook_req:
    add_list_item(w)

add_empty_line()

add_body_text(
    'Регистрация webhook выполняется методом setWebhook (HTTPS GET или POST):'
)
add_body_text(
    'https://api.telegram.org/bot<TOKEN>/setWebhook?url=https://<host>/<BOT_TOKEN>'
)

add_body_text(
    'Удаление webhook — методом deleteWebhook. Текущее состояние webhook можно '
    'получить через getWebhookInfo.'
)

add_empty_line()

# --- 3.4 Деплой и веб-приложение ---
add_body_text('3.4. Веб-приложение на Flask и развёртывание на pythonanywhere.com')
add_empty_line()

add_body_text(
    'Веб-приложение реализовано в модуле bot_webhook.py при помощи микрофреймворка '
    'Flask. Точкой входа служит WSGI-объект app = Flask(__name__). Приложение '
    'содержит два маршрута:'
)

routes = [
    'GET / — служебный health-check, возвращает строку «Lab4 webhook bot is running». '
    'Используется для проверки доступности приложения после деплоя;',
    'POST /<BOT_TOKEN> — основной эндпоинт, на который Telegram отправляет '
    'JSON-обновления. Использование токена бота в качестве пути URL — простая мера '
    'защиты: путь известен только владельцу бота и Telegram, поэтому посторонние '
    'не могут отправлять подложные Update.',
]
for r in routes:
    add_list_item(r)

add_empty_line()

add_body_text(
    'Последовательность развёртывания на сервисе pythonanywhere.com:'
)

deploy_steps = [
    'создать бесплатный аккаунт на pythonanywhere.com (поддерживается бесплатный '
    'тариф с одним веб-приложением);',
    'на вкладке Web → Add a new web app выбрать «Manual configuration» и версию '
    'Python 3.10/3.11;',
    'загрузить файлы лабораторной (bot_webhook.py, storage.py, ml_predict.py) и '
    'веса модели weights.pt из лаб. 3 в директорию /home/<user>/lab4/;',
    'в открывшемся WSGI-файле прописать: from bot_webhook import app as application;',
    'установить зависимости из requirements.txt командой '
    'pip install --user -r requirements.txt в Bash-консоли pythonanywhere;',
    'на вкладке Web → Environment variables прописать BOT_TOKEN со значением '
    'токена бота, полученного у @BotFather;',
    'нажать кнопку Reload и убедиться, что health-check отвечает по адресу '
    'https://<user>.pythonanywhere.com/.',
]
for s in deploy_steps:
    add_list_item(s)

add_empty_line()

add_body_text(
    'Telegram автоматически использует сертификат Let’s Encrypt, выданный для '
    'домена *.pythonanywhere.com, поэтому отдельная настройка TLS не требуется.'
)

add_empty_line()

# --- 3.5 setWebhook ---
add_body_text('3.5. Регистрация webhook’а')
add_empty_line()

add_body_text(
    'Для регистрации webhook реализован вспомогательный скрипт set_webhook.py. Он '
    'отправляет POST-запрос на метод setWebhook с URL веб-приложения и токеном '
    'бота в пути:'
)

add_body_text(
    'BOT_TOKEN=... python set_webhook.py set https://<user>.pythonanywhere.com'
)

add_body_text(
    'Telegram возвращает JSON-ответ вида {"ok": true, "result": true, '
    '"description": "Webhook was set"}. Состояние webhook’а можно проверить '
    'командой python set_webhook.py info, которая вызывает getWebhookInfo и '
    'отображает URL, число ожидающих обновлений (pending_update_count) и '
    'последнюю ошибку доставки (last_error_message).'
)

add_empty_line()

# --- 3.6 Формат сообщения ---
add_body_text('3.6. Формат сообщения, получаемого через webhook')
add_empty_line()

add_body_text(
    'После регистрации webhook в чат бота отправлено тестовое сообщение «hello». '
    'Telegram прислал POST-запрос на /<BOT_TOKEN> со следующим JSON-телом (фрагмент '
    'с подменённым chat_id):'
)

example_json = (
    '{\n'
    '  "update_id": 245932301,\n'
    '  "message": {\n'
    '    "message_id": 17,\n'
    '    "from": {\n'
    '      "id": 916411940,\n'
    '      "is_bot": false,\n'
    '      "first_name": "Vjatseslav",\n'
    '      "username": "TossSky",\n'
    '      "language_code": "ru"\n'
    '    },\n'
    '    "chat": {\n'
    '      "id": 916411940,\n'
    '      "first_name": "Vjatseslav",\n'
    '      "username": "TossSky",\n'
    '      "type": "private"\n'
    '    },\n'
    '    "date": 1747058201,\n'
    '    "text": "hello"\n'
    '  }\n'
    '}'
)
add_listing(example_json, 'Листинг 1 — Пример Update от Telegram через webhook')

add_empty_line()

add_body_text(
    'Ключевые поля объекта Update:'
)

fields = [
    'update_id — монотонно возрастающий идентификатор обновления, используется '
    'для дедупликации, если Telegram повторно отправит обновление;',
    'message — объект Message с информацией об отправленном сообщении: '
    'идентификатор сообщения, отправитель, чат, дата, текст или вложения;',
    'message.chat.id — идентификатор чата, нужен для отправки ответа через '
    'sendMessage;',
    'message.text — текст сообщения (для команд начинается с символа «/»);',
    'message.photo — массив объектов PhotoSize при отправке изображения; '
    'последний элемент массива содержит файл максимального разрешения и его '
    'file_id для загрузки через getFile.',
]
for f in fields:
    add_list_item(f)

add_empty_line()

# --- 3.7 Обработка сообщений и sendMessage ---
add_body_text('3.7. Обработка сообщений и отправка ответов через sendMessage')
add_empty_line()

add_body_text(
    'Обработка входящих обновлений реализована в функции webhook() (см. листинг 4): '
    'разбор JSON-тела, извлечение объекта message и диспатч по типу содержимого '
    '(команда, текст, фото). Отправка ответа выполняется в функции send_message — '
    'обычный HTTPS POST-запрос к методу sendMessage Telegram Bot API с обязательными '
    'параметрами chat_id и text. Полный URL:'
)

add_body_text(
    'https://api.telegram.org/bot<TOKEN>/sendMessage'
)

add_body_text(
    'При получении фотографии бот вызывает функцию download_photo: сначала '
    'обращается к getFile с file_id для получения относительного пути файла на '
    'серверах Telegram, затем скачивает файл по адресу '
    'https://api.telegram.org/file/bot<TOKEN>/<file_path> во временный файл, '
    'передаёт его в модель ResNet-18 (модуль ml_predict) и отправляет результат '
    'обратно в чат.'
)

add_empty_line()

# --- 3.8 Перенос функционала ---
add_body_text('3.8. Перенос функционала бота из лаб. 3')
add_empty_line()

add_body_text(
    'Все четыре команды, реализованные в лаб. 3, перенесены на webhook без '
    'изменения внешнего поведения:'
)

commands = [
    ('/register', 'регистрация пользователя. Сохраняет соль и хеш пароля в БД '
                  'через storage.register. Дублирующая регистрация запрещена.'),
    ('/login', 'аутентификация. Принимает пароль и сравнивает его хеш с сохранённым '
               'через storage.verify. При успехе устанавливает флаг logged_in=1.'),
    ('/predict', 'бинарная классификация изображения с помощью модели ResNet-18 из '
                 'лаб. 3. Скачивает фото из Telegram, прогоняет через модель '
                 '(ml_predict.predict_image), возвращает имя класса и вероятность.'),
    ('/logout', 'сброс флага logged_in. После выхода команда /predict снова '
                'требует аутентификации.'),
]
for cmd, desc in commands:
    add_bold_item(cmd, desc)

add_empty_line()

add_body_text(
    'Модель из лаб. 3 переиспользуется без переобучения. Адаптер ml_predict.py '
    'подключает к sys.path директорию Lab3, импортирует функции build_model и '
    'make_transforms, лениво загружает веса weights.pt при первом обращении и '
    'предоставляет единственную публичную функцию predict_image(path) → (label, prob).'
)

add_empty_line()

# --- 3.9 Сравнение polling и webhook ---
add_body_text('3.9. Сравнение реализаций polling и webhook')
add_empty_line()

add_body_text('Преимущества webhook:')

webhook_pros = [
    'мгновенная доставка обновлений (push-модель) без задержки, связанной с '
    'интервалом опроса. Latency определяется временем сетевого запроса от Telegram '
    'до сервера разработчика (порядка 100–300 мс);',
    'отсутствие постоянной исходящей нагрузки от приложения на API Telegram, '
    'что важно при ограничениях по сетевому трафику и для бесплатных тарифов '
    'облачных провайдеров;',
    'нативная масштабируемость через WSGI-сервер и балансировщик: один и тот же '
    'код принимает обновления параллельно на нескольких worker’ах;',
    'отсутствие необходимости держать процесс в активном цикле — приложение может '
    'быть serverless (FaaS), запускаясь только на время обработки запроса;',
    'возможность использовать бесплатные PaaS-сервисы (pythonanywhere.com, '
    'Vercel, Cloudflare Workers) без необходимости арендовать VPS.',
]
for p in webhook_pros:
    add_list_item(p)

add_empty_line()

add_body_text('Недостатки webhook:')

webhook_cons = [
    'обязательное наличие публичного HTTPS-эндпоинта с валидным TLS-сертификатом; '
    'это создаёт сложности при разработке (требуется ngrok или аналогичный '
    'туннель);',
    'ответ на запрос Telegram должен укладываться в 60 секунд, иначе обновление '
    'будет переотправлено. Это требует выноса долгих вычислений в фоновую очередь;',
    'усложнённая отладка: невозможно просто читать stdout — нужна работа с логами '
    'сервера и/или ngrok inspect;',
    'риск повторной обработки одного и того же обновления при ошибках '
    '(at-least-once-delivery), необходима идемпотентность по update_id;',
    'webhook требует открытого URL — это потенциальная поверхность атаки (DDoS, '
    'подделка обновлений), нужны меры защиты (секретный токен, валидация Update).',
]
for c in webhook_cons:
    add_list_item(c)

add_empty_line()

add_body_text('Преимущества polling:')

polling_pros = [
    'не требуется внешний URL и TLS — бот может работать из локальной сети, '
    'из-за NAT или с динамическим IP;',
    'простота разработки и отладки: достаточно запустить bot.infinity_polling() '
    'на ноутбуке;',
    'все ошибки видны в stdout процесса;',
    'отсутствует риск подделки запросов: бот сам опрашивает Telegram и доверяет '
    'только ему.',
]
for p in polling_pros:
    add_list_item(p)

add_empty_line()

add_body_text('Недостатки polling:')

polling_cons = [
    'задержка доставки обновлений (определяется интервалом long polling, '
    'обычно 25–30 секунд);',
    'постоянная исходящая нагрузка от приложения, бот должен непрерывно работать '
    'в активном цикле;',
    'однопоточность по умолчанию — при высокой нагрузке обновления обрабатываются '
    'последовательно;',
    'необходимость держать процесс «всегда живым» — нужны systemd, supervisor или '
    'аналогичный watchdog;',
    'плохо сочетается с serverless-средой и stateless-моделью;',
    'нельзя одновременно запускать два инстанса бота — Telegram возвращает ошибку '
    '«Conflict: terminated by other getUpdates request».',
]
for c in polling_cons:
    add_list_item(c)

add_empty_line()

# --- 3.10 Листинги ---
add_body_text(
    'В листингах 2–6 представлен исходный код всех модулей лабораторной работы.'
)
add_empty_line()

add_listing(read_code('storage.py'), 'Листинг 2 — Исходный код storage.py')
add_empty_line()
add_listing(read_code('bot_webhook.py'), 'Листинг 3 — Исходный код bot_webhook.py')
add_empty_line()
add_listing(read_code('ml_predict.py'), 'Листинг 4 — Исходный код ml_predict.py')
add_empty_line()
add_listing(read_code('set_webhook.py'), 'Листинг 5 — Исходный код set_webhook.py')
add_empty_line()
add_listing(read_code('bot_polling.py'), 'Листинг 6 — Исходный код bot_polling.py (для сравнения)')

doc.add_page_break()

# ===== 4. ВЫВОД =====
add_heading_text('4. Вывод')
add_empty_line()

add_body_text(
    'В ходе выполнения лабораторной работы реализован переход Telegram-бота из '
    'лаб. 3 с механизма long polling на webhook. Выявлены недостатки исходного '
    'способа хранения пользовательских данных (хранение в памяти процесса, '
    'plaintext-пароли, отсутствие потокобезопасности и персистентности), '
    'разработана и внедрена новая схема: SQLite-база с хешированием паролей '
    'алгоритмом PBKDF2-HMAC-SHA256, уникальной солью на каждого пользователя, '
    'константным сравнением хешей и блокировкой операций. Реализовано '
    'Flask-приложение, развёрнутое на pythonanywhere.com и принимающее '
    'POST-обновления от Telegram API; webhook зарегистрирован методом setWebhook. '
    'Все четыре команды (/register, /login, /predict, /logout) и интеграция с '
    'обученной в лаб. 3 моделью ResNet-18 сохранены без изменений в поведении. '
    'Выделены преимущества и недостатки обоих способов доставки обновлений: '
    'webhook предпочтителен для production-среды и масштабируемости, polling — '
    'для прототипирования и работы из-за NAT.'
)

# --- Сохранение ---
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Отчёт_ЛР4_Тоцкий_ВС.docx')
doc.save(output_path)
print(f'Report saved: {output_path}')
